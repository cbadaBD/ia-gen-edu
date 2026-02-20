import warnings
# Suprimir warnings de ScriptRunContext cuando se ejecuta en modo depuración
# Estos warnings son normales al usar debugger y no afectan la funcionalidad
warnings.filterwarnings('ignore', message='.*ScriptRunContext.*')
warnings.filterwarnings('ignore', message='.*missing ScriptRunContext.*')

import streamlit as st
import sys
import os
import re
from pathlib import Path
from datetime import datetime
from io import BytesIO

# Cargar variables de entorno desde .env si existe
try:
    from dotenv import load_dotenv
    load_dotenv()
except ImportError:
    pass

# Configurar AWS_REGION si no está definido (valor por defecto)
if 'AWS_REGION' not in os.environ:
    os.environ['AWS_REGION'] = 'us-east-1'

# Configurar página ANTES que cualquier otra cosa
st.set_page_config(page_title="Generador Educativo AI", page_icon="🤖", layout="wide")

# Inicializar variables de estado
DOCX_OK = False
SERVICES_OK = False

# Título principal
st.title("Generador de contenido educativo AI 🤖")
st.markdown("Genera material educativo con exportación a Word")

# Verificar imports paso a paso
with st.spinner("🔄 Verificando dependencias..."):
    # Agregar path
    try:
        sys.path.append(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
    except Exception as e:
        st.error(f"❌ Error agregando path: {e}")

    # Verificar python-docx
    try:
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.table import WD_ALIGN_VERTICAL
        from docx.shared import Inches, Pt, RGBColor
        DOCX_OK = True
    except ImportError as e:
        st.error(f"❌ python-docx no disponible: {e}")
        DOCX_OK = False

    # Verificar servicios Bedrock
    try:
        from core.bedrock_services import (
            generar_unidad_didactica,
            generar_sesion_aprendizaje,
            extraer_titulo_unidad_didactica,
            extraer_competencias_unidad_didactica,
            mejorar_documento_con_instruccion,
        )
        SERVICES_OK = True
    except Exception as e:
        st.error(f"❌ Error importando servicios: {e}")
        SERVICES_OK = False
    
    # Importar competencias (opcional, no crítico si falla)
    COMPETENCIAS_DISPONIBLES = False
    try:
        from core.competencias_curriculares import (
            obtener_todas_las_competencias,
            obtener_competencias_por_area,
            formatear_competencia_para_tabla,
            obtener_areas_curriculares_secundaria,
            obtener_grados_secundaria,
        )
        COMPETENCIAS_DISPONIBLES = True
    except Exception:
        # Si falla la importación, simplemente no mostrar el selector de competencias
        COMPETENCIAS_DISPONIBLES = False

# Fallback para menús de malla curricular si no se puede importar competencias_curriculares
AREAS_MALLA_FALLBACK = [
    "Desarrollo Personal, Ciudadanía y Cívica", "Ciencias Sociales", "Educación Física",
    "Arte y Cultura", "Comunicación", "Castellano como Segunda Lengua",
    "Inglés como Lengua Extranjera", "Matemática", "Ciencia y Tecnología",
    "Educación para el Trabajo", "Educación Religiosa",
]
GRADOS_MALLA_FALLBACK = ["1°", "2°", "3°", "4°", "5°"]

def dividir_contenido_largo_en_filas(item, contenido):
    """
    Divide el contenido largo de una celda en múltiples filas.
    Si el contenido tiene saltos de línea, cada línea adicional se convierte en una nueva fila
    con una celda vacía en la columna ITEM para mantener la alineación.
    
    Args:
        item: Texto del item (columna izquierda)
        contenido: Texto del contenido (columna derecha), puede tener saltos de línea
        
    Returns:
        Lista de filas de tabla en formato [item, contenido]
    """
    if not contenido:
        return [[item, ""]]
    
    # Dividir el contenido por saltos de línea
    lineas_contenido = contenido.split('\n')
    filas = []
    
    # Primera fila: item + primera línea de contenido
    if lineas_contenido:
        filas.append([item, lineas_contenido[0].strip()])
        
        # Filas adicionales: celda vacía + líneas restantes de contenido
        for linea_restante in lineas_contenido[1:]:
            if linea_restante.strip():  # Solo agregar si la línea no está vacía
                filas.append(["", linea_restante.strip()])  # Celda vacía en ITEM
    
    return filas if filas else [[item, ""]]


def normalizar_tabla_para_streamlit(contenido):
    """
    Normaliza el contenido de tabla para asegurar que siempre tenga formato ITEM | CONTENIDO
    y se muestre correctamente en Streamlit.
    Si encuentra contenido en la columna izquierda que no es un ITEM válido, lo mueve a la derecha.
    Si el contenido tiene saltos de línea, divide en múltiples filas con espacio en blanco en ITEM.
    
    Args:
        contenido: Contenido con tablas en formato markdown
        
    Returns:
        Contenido normalizado con tablas correctamente formateadas
    """
    if not contenido:
        return contenido
    
    lineas = contenido.split('\n')
    lineas_normalizadas = []
    dentro_tabla = False
    dentro_tabla_3_cols = False  # Para tabla de 3 columnas (Competencias transversales)
    tabla_headers = None  # Para tablas de 4+ columnas (Valores priorizados, Valores operativos, etc.)

    def es_fila_encabezado_multi_columna(fila_celdas):
        """Detecta si la fila es encabezado de tabla tipo Valores priorizados | Valores operativos | Enfoques transversales | Comportamientos observables."""
        if len(fila_celdas) < 3:
            return False
        texto_unido = ' '.join(fila_celdas).upper()
        palabras_clave = (
            'PRIORIZADOS', 'OPERATIVOS', 'TRANSVERSALES', 'OBSERVABLES',
            'VALORES PRIORIZADOS', 'ENFOQUES TRANSVERSALES', 'COMPORTAMIENTOS OBSERVABLES',
            'CARTA IDENTIDAD', 'IDENTIDAD ESA'
        )
        return any(p in texto_unido for p in palabras_clave)

    def es_item_valido(texto):
        """Determina si un texto es un ITEM válido"""
        if not texto or len(texto.strip()) == 0:
            return False
        # Si está vacío, no es un item válido
        if not texto or texto.strip() == "":
            return False
        # Si es muy largo, probablemente es contenido, no un item
        if len(texto) > 100:
            return False
        
        texto_upper = texto.upper().strip()
        texto_original = texto.strip()
        
        # Excluir frases que son claramente contenido, no items
        frases_contenido = [
            'MATERIALES PARA ESTUDIANTES',
            'MATERIALES PARA DOCENTE',
            'MATERIAL PARA ESTUDIANTES',
            'MATERIAL PARA DOCENTE',
            'PARA ESTUDIANTES',
            'PARA DOCENTE',
            'VALORES:',
            'ENFOQUES:',
            'COMPETENCIA:',
            'CAPACIDADES:',
            'DESEMPEÑOS:',
            'CRITERIOS:',
            'EVIDENCIAS:',
            'INSTRUMENTOS:',
            'RECURSOS:',
            'ACTIVIDADES:',
            'DIFICULTADES:',
            'MEJORAS:',
            'AJUSTES:'
        ]
        
        # Si contiene dos puntos y es una frase descriptiva, es contenido
        if ':' in texto_original and len(texto_original.split(':')) > 1:
            # Verificar si la parte antes de los dos puntos es una frase descriptiva
            parte_antes = texto_original.split(':')[0].strip().upper()
            if any(frase in parte_antes for frase in frases_contenido):
                return False
            # Si tiene más de 3 palabras antes de los dos puntos, probablemente es contenido
            if len(parte_antes.split()) > 3:
                return False
        
        # Si contiene "para" seguido de otra palabra, probablemente es contenido descriptivo
        if ' PARA ' in texto_upper or texto_upper.startswith('PARA '):
            return False
        
        palabras_item = ['TÍTULO', 'SITUACIÓN', 'COMPETENCIA', 'CAPACIDAD', 
                        'EVIDENCIA', 'INSTRUMENTO', 'VALOR', 'SECUENCIA', 
                        'ENFOQUE', 'SESIÓN', 'MATERIAL', 'REFLEXIÓN', 'ESTÁNDAR',
                        'DESEMPEÑO', 'PROPÓSITO', 'ORGANIZACIÓN', 'EVALUACIÓN',
                        'DATOS', 'CRITERIO', 'MOMENTO', 'DIDÁCTICA', 'INFORMATIVOS',
                        'SIGNIFICATIVA', 'PRECISADOS', 'APRENDIZAJE',
                        'DEFINICIÓN CONCEPTUAL', 'VALORES PRIORIZADOS', 'VALORES OPERATIVOS',
                        'ENFOQUES TRANSVERSALES', 'COMPORTAMIENTOS OBSERVABLES']
        
        # Limpiar formato markdown bold para análisis
        texto_sin_bold = texto_original.replace('**', '').strip()
        texto_sin_bold_upper = texto_sin_bold.upper()
        
        # Solo considerar como item si:
        # 1. Es muy corto y está en mayúsculas (típico de encabezados)
        # 2. Empieza con ** (formato markdown bold) - estos son siempre items
        # 3. Es una palabra clave específica Y no es una frase descriptiva
        es_palabra_clave = any(palabra in texto_sin_bold_upper for palabra in palabras_item)
        
        # Si empieza con **, es definitivamente un item (formato markdown bold)
        if texto_original.strip().startswith('**') and texto_original.strip().endswith('**'):
            # Verificar que no sea una frase de contenido excluida
            texto_limpio = texto_sin_bold_upper
            if not any(frase in texto_limpio for frase in ['MATERIALES PARA', 'PARA ESTUDIANTES', 'PARA DOCENTE']):
                return True
        
        # Si es una palabra clave pero es una frase descriptiva, no es un item
        if es_palabra_clave:
            # Verificar si es solo la palabra clave o una frase
            palabras_texto = texto_sin_bold_upper.split()
            # Si tiene más de 2 palabras y contiene "PARA", es contenido
            if len(palabras_texto) > 2 and 'PARA' in palabras_texto:
                return False
            # Si tiene más de 5 palabras en total y NO está en negrita, probablemente es contenido
            if len(palabras_texto) > 5 and not texto_original.strip().startswith('**'):
                return False
        
        return (
            (len(texto_sin_bold) < 50 and texto_sin_bold.isupper() and len(texto_sin_bold.split()) <= 5) or
            (texto_original.strip().startswith('**') and len(texto_sin_bold.split()) <= 6) or
            (es_palabra_clave and len(texto_sin_bold.split()) <= 5 and not any(frase in texto_sin_bold_upper for frase in ['MATERIALES PARA', 'PARA ESTUDIANTES', 'PARA DOCENTE']))
        )
    
    def obtener_ultima_fila_info():
        """Obtiene información de la última fila de tabla"""
        if not lineas_normalizadas or not lineas_normalizadas[-1].startswith('|'):
            return None, None, None
        ultima = lineas_normalizadas[-1]
        partes = ultima.split('|')
        if len(partes) >= 3:
            item = partes[1].strip()
            contenido = partes[2].strip()
            return item, contenido, ultima
        elif len(partes) >= 2:
            item = partes[1].strip()
            return item, "", ultima
        return None, None, None
    
    i = 0
    while i < len(lineas):
        linea = lineas[i]
        linea_stripped = linea.strip()
        
        # Detectar inicio de tabla
        if re.match(r'^\s*\|.*\|\s*$', linea_stripped) and linea_stripped.count('|') >= 2:
            es_separador = re.match(r'^\s*\|[\s\-\:]+\|\s*$', linea_stripped)
            
            # Si hay una línea vacía antes, terminar tabla anterior
            if i > 0 and not lineas[i-1].strip() and (dentro_tabla or dentro_tabla_3_cols):
                dentro_tabla = False
                dentro_tabla_3_cols = False
                tabla_headers = None
            
            # Detectar si es tabla de 3 columnas (Competencias transversales)
            partes_temp = linea_stripped.split('|')
            fila_temp = [celda.strip() for celda in partes_temp if celda.strip()]
            es_tabla_3_cols = (
                len(fila_temp) == 3 and 
                not es_separador and
                any(palabra in ' '.join(fila_temp).upper() for palabra in ['COMPETENCIAS TRANSVERSALES', 'ESTÁNDARES', 'INSTRUMENTO'])
            )
            
            if es_tabla_3_cols and not dentro_tabla_3_cols:
                # Iniciar tabla de 3 columnas
                dentro_tabla_3_cols = True
                dentro_tabla = False
                tabla_headers = None
                lineas_normalizadas.append(linea_stripped)
            elif dentro_tabla_3_cols:
                # Continuar tabla de 3 columnas - mantener formato original
                if es_separador:
                    lineas_normalizadas.append(linea_stripped)
                else:
                    # Verificar si sigue siendo tabla de 3 columnas
                    partes = linea.split('|')
                    fila = [celda.strip() for celda in partes]
                    while fila and not fila[0]:
                        fila.pop(0)
                    while fila and not fila[-1]:
                        fila.pop()
                    
                    if len(fila) == 3:
                        # Mantener formato de 3 columnas
                        lineas_normalizadas.append(linea_stripped)
                    else:
                        # Ya no es tabla de 3 columnas, cambiar modo
                        dentro_tabla_3_cols = False
                        dentro_tabla = True
                        # Procesar como tabla normal
                        if len(fila) >= 2:
                            primera_col = fila[0]
                            resto_contenido = ' '.join([c for c in fila[1:] if c])
                            es_item = es_item_valido(primera_col)
                            if not es_item:
                                contenido_completo = primera_col + (' ' + resto_contenido if resto_contenido else '')
                                filas = dividir_contenido_largo_en_filas("", contenido_completo)
                                for fila_item, fila_contenido in filas:
                                    lineas_normalizadas.append(f"| {fila_item} | {fila_contenido} |")
                            else:
                                filas = dividir_contenido_largo_en_filas(primera_col, resto_contenido)
                                for fila_item, fila_contenido in filas:
                                    lineas_normalizadas.append(f"| {fila_item} | {fila_contenido} |")
            else:
                # Tabla normal de 2 columnas
                dentro_tabla = True
                dentro_tabla_3_cols = False
                
                if es_separador:
                    # Línea separadora, mantenerla
                    lineas_normalizadas.append(linea_stripped)
                else:
                    # Fila de tabla: normalizar a formato ITEM | CONTENIDO
                    partes = linea.split('|')
                    fila = [celda.strip() for celda in partes]
                    # Eliminar celdas vacías al inicio y final
                    while fila and not fila[0]:
                        fila.pop(0)
                    while fila and not fila[-1]:
                        fila.pop()
                    
                    # Normalizar a 2 columnas
                    if len(fila) == 0:
                        pass  # Fila vacía, saltar
                    elif len(fila) == 1:
                        # Una columna: determinar si es ITEM o CONTENIDO
                        contenido_unico = fila[0]
                        es_item = es_item_valido(contenido_unico)
                        
                        # PRIMERO: Verificar si la última fila tiene ITEM sin CONTENIDO
                        item_ultimo, contenido_ultimo, ultima_linea = obtener_ultima_fila_info()
                        if item_ultimo and item_ultimo != "" and (not contenido_ultimo or len(contenido_ultimo) < 30):
                            # La última fila tiene ITEM sin CONTENIDO, agregar este contenido ahí
                            if contenido_ultimo:
                                contenido_combinado = contenido_ultimo + '\n' + contenido_unico
                            else:
                                contenido_combinado = contenido_unico
                            # Dividir contenido largo en múltiples filas
                            filas = dividir_contenido_largo_en_filas(item_ultimo, contenido_combinado)
                            # Reemplazar la última fila con la primera fila dividida
                            if filas:
                                lineas_normalizadas[-1] = f"| {filas[0][0]} | {filas[0][1]} |"
                                # Agregar filas adicionales con espacio en blanco en ITEM
                                for fila_item, fila_contenido in filas[1:]:
                                    lineas_normalizadas.append(f"| {fila_item} | {fila_contenido} |")
                        elif es_item:
                            # Es un ITEM nuevo
                            lineas_normalizadas.append(f"| {contenido_unico} | |")
                        else:
                            # Es CONTENIDO pero no hay ITEM previo sin CONTENIDO
                            # Dividir contenido largo en múltiples filas con celda vacía en ITEM
                            filas = dividir_contenido_largo_en_filas("", contenido_unico)
                            for fila_item, fila_contenido in filas:
                                lineas_normalizadas.append(f"| {fila_item} | {fila_contenido} |")
                    elif len(fila) == 2:
                        # Dos columnas: formato ITEM | CONTENIDO
                        item = fila[0]
                        contenido = fila[1]
                        filas = dividir_contenido_largo_en_filas(item, contenido)
                        for fila_item, fila_contenido in filas:
                            lineas_normalizadas.append(f"| {fila_item} | {fila_contenido} |")
                    elif len(fila) >= 4:
                        # Tabla con 4+ columnas (Valores priorizados | Valores operativos | Enfoques transversales | Comportamientos observables)
                        if len(fila) > 2 and es_fila_encabezado_multi_columna(fila) and tabla_headers is None:
                            tabla_headers = [c.strip() for c in fila]
                            # No añadir esta fila como datos; se usará como nombres de ítem para las siguientes filas
                        elif tabla_headers is not None and len(fila) == len(tabla_headers):
                            # Fila de datos: expandir en una fila ITEM|CONTENIDO por cada columna
                            for j in range(len(fila)):
                                item_nombre = tabla_headers[j]
                                contenido_celda = fila[j].strip()
                                filas = dividir_contenido_largo_en_filas(item_nombre, contenido_celda)
                                for fila_item, fila_contenido in filas:
                                    lineas_normalizadas.append(f"| {fila_item} | {fila_contenido} |")
                        else:
                            if tabla_headers is not None:
                                tabla_headers = None
                            # Dos o más columnas: verificar si la primera es realmente un ITEM
                            primera_col = fila[0]
                            resto_contenido = ' '.join([c for c in fila[1:] if c])

                            # Limpiar formato markdown bold del item si está presente
                            item_limpio = primera_col.strip()
                            if item_limpio.startswith('**') and item_limpio.endswith('**'):
                                item_limpio = item_limpio[2:-2].strip()

                            # Verificar si es un item válido (usar el texto limpio para verificación)
                            es_item = es_item_valido(primera_col)

                            # Si la primera columna NO es un ITEM válido, mover todo a la derecha
                            if not es_item:
                                # La primera columna es contenido, mover todo a la derecha con celda vacía
                                contenido_completo = primera_col
                                if resto_contenido:
                                    contenido_completo = primera_col + ' ' + resto_contenido
                                # Dividir contenido largo en múltiples filas
                                filas = dividir_contenido_largo_en_filas("", contenido_completo)
                                for fila_item, fila_contenido in filas:
                                    lineas_normalizadas.append(f"| {fila_item} | {fila_contenido} |")
                            else:
                                # La primera columna es un ITEM válido
                                # Usar el item original (con ** si estaba) para mantener formato
                                # Dividir contenido largo en múltiples filas si tiene saltos de línea
                                filas = dividir_contenido_largo_en_filas(primera_col, resto_contenido)
                                for fila_item, fila_contenido in filas:
                                    lineas_normalizadas.append(f"| {fila_item} | {fila_contenido} |")
        else:
            # Línea fuera de tabla (sin formato |)
            # Si hay línea vacía, terminar la tabla actual (ya sea de 2 o 3 columnas)
            if not linea_stripped:
                # Línea vacía: terminar tabla actual
                if dentro_tabla_3_cols:
                    dentro_tabla_3_cols = False
                if dentro_tabla:
                    dentro_tabla = False
                tabla_headers = None
                lineas_normalizadas.append(linea)
            elif dentro_tabla_3_cols and linea_stripped:
                # Contenido fuera de tabla de 3 columnas: terminar tabla
                dentro_tabla_3_cols = False
                lineas_normalizadas.append(linea)
            elif dentro_tabla and linea_stripped:
                # Si estamos dentro de una tabla y encontramos contenido sin |,
                # SIEMPRE agregarlo como nueva fila con espacio en blanco en ITEM
                if (lineas_normalizadas and 
                    lineas_normalizadas[-1].startswith('|') and
                    not linea_stripped.startswith('#') and
                    not re.match(r'^\s*\|[\s\-\:]+\|\s*$', linea_stripped)):
                    item_ultimo, contenido_ultimo, ultima_linea = obtener_ultima_fila_info()
                    if item_ultimo is not None and item_ultimo != "":
                        # Agregar como nueva fila con celda vacía en ITEM
                        # Dividir en múltiples filas si es necesario
                        filas = dividir_contenido_largo_en_filas("", linea_stripped)
                        for fila_item, fila_contenido in filas:
                            lineas_normalizadas.append(f"| {fila_item} | {fila_contenido} |")
                    else:
                        dentro_tabla = False
                        dentro_tabla_3_cols = False
                        tabla_headers = None
                        lineas_normalizadas.append(linea)
                else:
                    dentro_tabla = False
                    dentro_tabla_3_cols = False
                    tabla_headers = None
                    lineas_normalizadas.append(linea)
            else:
                lineas_normalizadas.append(linea)
        
        i += 1
    
    # Post-procesamiento: dividir contenido largo en las filas finales y corregir items mal ubicados
    lineas_finales = []
    dentro_tabla_3_cols_post = False
    for linea in lineas_normalizadas:
        if '|' in linea and not re.match(r'^\s*\|[\s\-\:]+\|\s*$', linea.strip()):
            # Verificar si es tabla de 3 columnas
            partes_temp = linea.split('|')
            fila_temp = [celda.strip() for celda in partes_temp if celda.strip()]
            es_tabla_3_cols_linea = (
                len(fila_temp) == 3 and
                any(palabra in ' '.join(fila_temp).upper() for palabra in ['COMPETENCIAS TRANSVERSALES', 'ESTÁNDARES', 'INSTRUMENTO', 'SE DESENVUELVE', 'GESTIONA'])
            )
            
            if es_tabla_3_cols_linea or (dentro_tabla_3_cols_post and len(fila_temp) == 3):
                dentro_tabla_3_cols_post = True
                # Mantener tabla de 3 columnas sin modificar
                lineas_finales.append(linea)
                continue
            elif dentro_tabla_3_cols_post and len(fila_temp) != 3:
                # Terminar tabla de 3 columnas
                dentro_tabla_3_cols_post = False
                # Continuar procesando como tabla normal
            
            # Procesar como tabla normal (2 columnas)
            partes = linea.split('|')
            if len(partes) >= 3:
                item = partes[1].strip()
                contenido_celda = partes[2].strip()
                
                # Limpiar formato markdown bold del item para análisis
                item_limpio = item.replace('**', '').strip() if item else ""
                item_upper = item_limpio.upper() if item_limpio else ""
                
                # Detectar frases que son claramente contenido, no items
                frases_contenido_detectadas = [
                    'MATERIALES PARA ESTUDIANTES',
                    'MATERIALES PARA DOCENTE',
                    'MATERIAL PARA ESTUDIANTES',
                    'MATERIAL PARA DOCENTE',
                    'PARA ESTUDIANTES',
                    'PARA DOCENTE',
                    'VALORES:',
                    'ENFOQUES:',
                    'COMPETENCIA:',
                    'CAPACIDADES:',
                    'DESEMPEÑOS:',
                    'CRITERIOS:',
                    'EVIDENCIAS:',
                    'INSTRUMENTOS:',
                    'RECURSOS:',
                    'ACTIVIDADES:'
                ]
                
                es_frase_contenido = any(frase in item_upper for frase in frases_contenido_detectadas)
                
                # Si el item contiene ":" y es una frase descriptiva, es contenido
                if item_limpio and ':' in item_limpio:
                    parte_antes = item_limpio.split(':')[0].strip().upper()
                    if any(frase in parte_antes for frase in frases_contenido_detectadas):
                        es_frase_contenido = True
                    # Si tiene más de 3 palabras antes de los dos puntos, probablemente es contenido
                    if len(parte_antes.split()) > 3:
                        es_frase_contenido = True
                
                # Si el item está en negrita (**), es definitivamente un item válido
                es_item_en_negrita = item and item.strip().startswith('**') and item.strip().endswith('**')
                
                # Si el item no es válido O es una frase de contenido (y NO está en negrita), mover todo a la derecha
                if item and (not es_item_valido(item) or es_frase_contenido) and not es_item_en_negrita:
                    # El item es en realidad contenido, mover todo a la derecha
                    if contenido_celda:
                        contenido_completo = item + ' ' + contenido_celda
                    else:
                        contenido_completo = item
                    filas = dividir_contenido_largo_en_filas("", contenido_completo)
                    for fila_item, fila_contenido in filas:
                        lineas_finales.append(f"| {fila_item} | {fila_contenido} |")
                # Si el item está en negrita, mantenerlo en la izquierda y el contenido en la derecha
                elif es_item_en_negrita:
                    # Asegurar que el contenido esté en la columna derecha
                    # Si el contenido está vacío o es muy corto, puede que esté mezclado con el item
                    if not contenido_celda or len(contenido_celda) < 10:
                        # El contenido puede estar en la misma celda que el item, verificar
                        item_limpio = item.replace('**', '').strip()
                        # Si el item tiene contenido después de los **, separarlo
                        if '**' in item and len(item.split('**')) > 2:
                            partes_item = item.split('**')
                            if len(partes_item) >= 3:
                                item_final = '**' + partes_item[1] + '**'
                                contenido_restante = ' '.join(partes_item[2:]).strip()
                                if contenido_restante:
                                    contenido_celda = contenido_restante + (' ' + contenido_celda if contenido_celda else '')
                                    item = item_final
                    
                    # Dividir contenido largo en múltiples filas si es necesario
                    if '\n' in contenido_celda or len(contenido_celda) > 200:
                        filas = dividir_contenido_largo_en_filas(item, contenido_celda)
                        for fila_item, fila_contenido in filas:
                            lineas_finales.append(f"| {fila_item} | {fila_contenido} |")
                    else:
                        lineas_finales.append(f"| {item} | {contenido_celda} |")
                # Si el contenido tiene saltos de línea, dividir en múltiples filas
                elif '\n' in contenido_celda or (len(contenido_celda) > 200 and item):
                    filas = dividir_contenido_largo_en_filas(item, contenido_celda)
                    for fila_item, fila_contenido in filas:
                        lineas_finales.append(f"| {fila_item} | {fila_contenido} |")
                else:
                    lineas_finales.append(linea)
            else:
                lineas_finales.append(linea)
        else:
            lineas_finales.append(linea)
    
    resultado = '\n'.join(lineas_finales)
    
    # Asegurar que las tablas tengan el formato correcto para Streamlit
    # Streamlit requiere una línea separadora después del encabezado
    lineas_resultado = resultado.split('\n')
    lineas_formateadas = []
    dentro_tabla = False
    ultima_fila_era_encabezado = False
    
    i = 0
    while i < len(lineas_resultado):
        linea = lineas_resultado[i]
        linea_stripped = linea.strip()
        
        # Detectar si es una línea de tabla
        if '|' in linea_stripped and linea_stripped.count('|') >= 2:
            # Verificar si es un separador
            es_separador = re.match(r'^\s*\|[\s\-\:]+\|\s*$', linea_stripped)
            
            if es_separador:
                # Ya hay un separador, mantenerlo pero asegurar formato correcto
                num_cols = linea_stripped.count('|') - 1
                separador = '|' + '|'.join(['---' for _ in range(num_cols)]) + '|'
                lineas_formateadas.append(separador)
                dentro_tabla = True
                ultima_fila_era_encabezado = False
            else:
                # Es una fila de datos o encabezado
                # Verificar si es encabezado (contiene ITEM y CONTENIDO)
                es_encabezado = ('ITEM' in linea_stripped.upper() and 'CONTENIDO' in linea_stripped.upper())
                
                # Si la última fila era encabezado y no había separador, agregarlo ahora
                if ultima_fila_era_encabezado:
                    num_cols = lineas_formateadas[-1].count('|') - 1
                    separador = '|' + '|'.join(['---' for _ in range(num_cols)]) + '|'
                    lineas_formateadas.append(separador)
                    ultima_fila_era_encabezado = False
                
                lineas_formateadas.append(linea_stripped)
                dentro_tabla = True
                
                # Si es encabezado, marcar para agregar separador después
                if es_encabezado:
                    # Verificar si la siguiente línea es un separador
                    if i + 1 < len(lineas_resultado):
                        siguiente = lineas_resultado[i + 1].strip()
                        es_sig_separador = re.match(r'^\s*\|[\s\-\:]+\|\s*$', siguiente)
                        if not es_sig_separador:
                            ultima_fila_era_encabezado = True
                    else:
                        # Es la última línea y es encabezado, agregar separador
                        num_cols = linea_stripped.count('|') - 1
                        separador = '|' + '|'.join(['---' for _ in range(num_cols)]) + '|'
                        lineas_formateadas.append(separador)
        else:
            # Si salimos de una tabla y la última fila era encabezado, agregar separador
            if dentro_tabla and ultima_fila_era_encabezado:
                num_cols = lineas_formateadas[-1].count('|') - 1
                separador = '|' + '|'.join(['---' for _ in range(num_cols)]) + '|'
                lineas_formateadas.append(separador)
                ultima_fila_era_encabezado = False
            
            dentro_tabla = False
            if linea_stripped:
                lineas_formateadas.append(linea)
            elif lineas_formateadas and lineas_formateadas[-1].strip():
                lineas_formateadas.append('')
        
        i += 1
    
    # Si terminamos dentro de una tabla y la última fila era encabezado, agregar separador
    if dentro_tabla and ultima_fila_era_encabezado:
        num_cols = lineas_formateadas[-1].count('|') - 1
        separador = '|' + '|'.join(['---' for _ in range(num_cols)]) + '|'
        lineas_formateadas.append(separador)
    
    return '\n'.join(lineas_formateadas)

def convertir_tablas_markdown_a_html(contenido_markdown):
    """
    Convierte tablas markdown a HTML con los mismos estilos que la tabla de DATOS INFORMATIVOS.
    """
    lineas = contenido_markdown.split('\n')
    resultado_html = []
    dentro_tabla = False
    filas_tabla = []
    es_tabla_3_cols = False
    
    i = 0
    while i < len(lineas):
        linea = lineas[i]
        linea_stripped = linea.strip()
        
        # Detectar inicio de tabla markdown
        if '|' in linea_stripped and linea_stripped.count('|') >= 2:
            es_separador = re.match(r'^\s*\|[\s\-\:]+\|\s*$', linea_stripped)
            
            if es_separador:
                # Línea separadora: continuar procesando la tabla
                dentro_tabla = True
                i += 1
                continue
            
            # Verificar si es tabla de 3 columnas
            partes_temp = linea_stripped.split('|')
            fila_temp = [celda.strip() for celda in partes_temp if celda.strip()]
            es_tabla_3_cols = (
                len(fila_temp) == 3 and
                any(palabra in ' '.join(fila_temp).upper() for palabra in ['COMPETENCIAS TRANSVERSALES', 'ESTÁNDARES', 'INSTRUMENTO'])
            )
            
            if not dentro_tabla:
                # Iniciar nueva tabla
                dentro_tabla = True
                filas_tabla = []
            
            # Procesar fila de tabla
            partes = linea_stripped.split('|')
            # Preservar todas las celdas, incluyendo las vacías intencionales
            fila = [celda.strip() for celda in partes]
            # Solo eliminar celdas vacías al inicio y final si están completamente vacías (sin espacios)
            # Pero preservar celdas que son intencionalmente vacías entre columnas
            if len(fila) > 2:
                # Hay al menos 3 partes (inicio vacío, columnas, final vacío)
                # Mantener solo las columnas del medio
                fila = fila[1:-1] if not fila[0] and not fila[-1] else fila
            elif len(fila) == 2:
                # Solo 2 partes: podría ser inicio vacío + columna o columna + final vacío
                if not fila[0] and fila[1]:
                    fila = ["", fila[1]]
                elif fila[0] and not fila[1]:
                    fila = [fila[0], ""]
                else:
                    fila = fila
            
            # Asegurar que siempre tenga al menos 2 columnas para tablas de 2 columnas
            if fila:
                # Si es tabla de 2 columnas y solo tiene 1 columna, agregar columna vacía
                if not es_tabla_3_cols and len(fila) == 1:
                    # Determinar si es ITEM o CONTENIDO basándose en el contenido
                    contenido_unico = fila[0]
                    # Si la última fila tenía ITEM sin CONTENIDO, este es CONTENIDO
                    if filas_tabla and len(filas_tabla[-1]) == 1:
                        # La última fila solo tiene ITEM, agregar este como CONTENIDO
                        filas_tabla[-1].append(contenido_unico)
                    else:
                        # Es un nuevo ITEM o CONTENIDO, agregar como fila con 2 columnas
                        # Si parece ser un ITEM (corto, en mayúsculas, o con **), poner en primera columna
                        es_item = (
                            contenido_unico.startswith('**') or
                            (len(contenido_unico) < 100 and contenido_unico.isupper()) or
                            any(palabra in contenido_unico.upper() for palabra in ['COMPETENCIAS:', 'CAPACIDADES:', 'CRITERIOS:', 'CONTENIDOS:', 'EVIDENCIA:', 'INSTRUMENTO:'])
                        )
                        if es_item:
                            filas_tabla.append([contenido_unico, ""])
                        else:
                            # Es CONTENIDO, agregar a la última fila si tenía ITEM sin CONTENIDO
                            if filas_tabla and len(filas_tabla[-1]) >= 1 and not filas_tabla[-1][-1]:
                                filas_tabla[-1][-1] = contenido_unico
                            else:
                                filas_tabla.append(["", contenido_unico])
                elif len(fila) >= 2:
                    # Ya tiene 2 o más columnas, agregar tal cual
                    filas_tabla.append(fila[:2] if not es_tabla_3_cols else fila[:3])
                else:
                    filas_tabla.append(fila)
        else:
            # Línea fuera de tabla
            if dentro_tabla and filas_tabla:
                # Convertir tabla acumulada a HTML
                # Determinar número de columnas basándose en todas las filas
                num_cols = 2
                if filas_tabla:
                    # Verificar si es tabla de 3 columnas
                    primera_fila_cols = len(filas_tabla[0])
                    todas_3_cols = all(len(f) == 3 for f in filas_tabla if f)
                    es_tabla_3_cols_final = (
                        todas_3_cols and 
                        primera_fila_cols == 3 and
                        any(palabra in ' '.join(filas_tabla[0]).upper() for palabra in ['COMPETENCIAS TRANSVERSALES', 'ESTÁNDARES', 'INSTRUMENTO'])
                    )
                    if es_tabla_3_cols_final:
                        num_cols = 3
                    else:
                        # Para tablas de 2 columnas, asegurar que todas tengan 2 columnas
                        num_cols = 2
                        # Normalizar todas las filas a 2 columnas
                        filas_normalizadas = []
                        for fila in filas_tabla:
                            if len(fila) == 0:
                                filas_normalizadas.append(["", ""])
                            elif len(fila) == 1:
                                # Si solo tiene 1 columna, determinar si es ITEM o CONTENIDO
                                contenido_unico = fila[0]
                                # Si la última fila normalizada tenía ITEM sin CONTENIDO, agregar aquí
                                if filas_normalizadas and filas_normalizadas[-1][0] and not filas_normalizadas[-1][1]:
                                    filas_normalizadas[-1][1] = contenido_unico
                                elif filas_normalizadas and not filas_normalizadas[-1][0] and filas_normalizadas[-1][1]:
                                    # La última fila tenía CONTENIDO sin ITEM, este es continuación del CONTENIDO
                                    filas_normalizadas[-1][1] += " " + contenido_unico
                                else:
                                    # Es un nuevo ITEM o CONTENIDO
                                    es_item = (
                                        contenido_unico.startswith('**') or
                                        (len(contenido_unico) < 100 and contenido_unico.isupper() and not contenido_unico.startswith('---')) or
                                        any(palabra in contenido_unico.upper() for palabra in ['COMPETENCIAS:', 'CAPACIDADES:', 'CRITERIOS:', 'CONTENIDOS:', 'EVIDENCIA:', 'INSTRUMENTO:'])
                                    )
                                    if es_item:
                                        filas_normalizadas.append([contenido_unico, ""])
                                    else:
                                        filas_normalizadas.append(["", contenido_unico])
                            else:
                                # Ya tiene 2 o más columnas, tomar solo las primeras 2
                                item_val = fila[0] if len(fila) > 0 else ""
                                contenido_val = fila[1] if len(fila) > 1 else ""
                                filas_normalizadas.append([item_val, contenido_val])
                        filas_tabla = filas_normalizadas
                
                html_tabla = '<table style="width: 100%; border-collapse: collapse;">\n'
                
                for idx_fila, fila in enumerate(filas_tabla):
                    html_tabla += '<tr>\n'
                    
                    # Determinar si es encabezado (primera fila con palabras clave)
                    es_encabezado = (
                        idx_fila == 0 and 
                        any(palabra in ' '.join(fila).upper() for palabra in ['ITEM', 'CONTENIDO', 'COMPETENCIAS', 'ESTÁNDARES', 'INSTRUMENTO'])
                    )
                    
                    # Asegurar que la fila tenga exactamente num_cols columnas
                    while len(fila) < num_cols:
                        fila.append("")
                    fila = fila[:num_cols]
                    
                    for idx_col, celda in enumerate(fila):
                        # Estilo para celdas de encabezado (colores para modo oscuro)
                        if es_encabezado:
                            estilo = 'background-color: #2C3E50; color: #FFFFFF; padding: 8px; border: 1px solid #555; font-weight: bold;'
                        else:
                            # Estilo para celdas de datos (colores para modo oscuro)
                            estilo = 'background-color: #1E1E1E; color: #E0E0E0; padding: 8px; border: 1px solid #555;'
                        
                        html_tabla += f'<td style="{estilo}">{celda}</td>\n'
                    
                    html_tabla += '</tr>\n'
                
                html_tabla += '</table>\n'
                resultado_html.append(html_tabla)
                
                # Resetear estado
                dentro_tabla = False
                filas_tabla = []
                es_tabla_3_cols = False
            
            # Agregar línea fuera de tabla
            resultado_html.append(linea)
        
        i += 1
    
    # Procesar última tabla si termina el contenido
    if dentro_tabla and filas_tabla:
        num_cols = len(filas_tabla[0]) if filas_tabla else 2
        html_tabla = '<table style="width: 100%; border-collapse: collapse;">\n'
        
        for idx_fila, fila in enumerate(filas_tabla):
            html_tabla += '<tr>\n'
            
            es_encabezado = (
                idx_fila == 0 and 
                any(palabra in ' '.join(fila).upper() for palabra in ['ITEM', 'CONTENIDO', 'COMPETENCIAS', 'ESTÁNDARES', 'INSTRUMENTO'])
            )
            
            for celda in fila:
                # Estilo para modo oscuro
                if es_encabezado:
                    estilo = 'background-color: #2C3E50; color: #FFFFFF; padding: 8px; border: 1px solid #555; font-weight: bold;'
                else:
                    estilo = 'background-color: #1E1E1E; color: #E0E0E0; padding: 8px; border: 1px solid #555;'
                html_tabla += f'<td style="{estilo}">{celda}</td>\n'
            
            html_tabla += '</tr>\n'
        
        html_tabla += '</table>\n'
        resultado_html.append(html_tabla)
    
    return '\n'.join(resultado_html)

def generar_tabla_secuencia_sesiones(num_sesiones, duracion_horas=4, sesiones_data=None):
    """
    Genera la tabla dinámica de secuencia de sesiones.
    
    Args:
        num_sesiones: Número total de sesiones
        duracion_horas: Duración de cada sesión en horas (por defecto 4)
        sesiones_data: Lista opcional de dicts con keys titulo, criterio, actividades (generados por IA)
    
    Returns:
        HTML de la tabla de secuencia de sesiones
    """
    if sesiones_data is None:
        sesiones_data = []
    # Asegurar que hay un dict por cada sesión
    while len(sesiones_data) < num_sesiones:
        sesiones_data.append({"titulo": "", "criterio": "", "actividades": ""})
    sesiones_data = sesiones_data[:num_sesiones]
    
    def escapar_html(s):
        if not s:
            return ""
        return str(s).replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;").replace('"', "&quot;")
    
    html_tabla = '<h3>IV. SECUENCIA DE SESIONES</h3>\n'
    html_tabla += '<table style="width: 100%; border-collapse: collapse;">\n'
    
    # Organizar sesiones en filas de 2 columnas
    sesiones_por_fila = 2
    num_filas = (num_sesiones + sesiones_por_fila - 1) // sesiones_por_fila  # Redondeo hacia arriba
    
    for fila_idx in range(num_filas):
        html_tabla += '<tr>\n'
        
        for col_idx in range(sesiones_por_fila):
            sesion_num = fila_idx * sesiones_por_fila + col_idx + 1
            
            if sesion_num > num_sesiones:
                # Si no hay más sesiones, dejar celda vacía
                html_tabla += '<td style="padding: 0; border: none;"></td>\n'
            else:
                datos = sesiones_data[sesion_num - 1] if sesion_num <= len(sesiones_data) else {}
                titulo = datos.get("titulo", "")
                criterio = datos.get("criterio", "")
                actividades = datos.get("actividades", "")
                # Crear celda para esta sesión
                html_tabla += '<td style="padding: 5px; border: 1px solid #555; vertical-align: top;">\n'
                
                # Encabezado de sesión
                html_tabla += f'<table style="width: 100%; border-collapse: collapse;">\n'
                html_tabla += '<tr>\n'
                html_tabla += f'<td colspan="2" style="background-color: #2C3E50; color: #FFFFFF; padding: 6px; border: 1px solid #555; font-weight: bold;">Sesión N° {sesion_num}: <span style="float: right;">({duracion_horas} horas)</span></td>\n'
                html_tabla += '</tr>\n'
                
                # Fila de Título
                html_tabla += '<tr>\n'
                html_tabla += '<td style="background-color: #1E1E1E; color: #E0E0E0; padding: 6px; border: 1px solid #555; font-weight: bold; width: 30%;">Título:</td>\n'
                html_tabla += f'<td style="background-color: #1E1E1E; color: #E0E0E0; padding: 6px; border: 1px solid #555;">{escapar_html(titulo)}</td>\n'
                html_tabla += '</tr>\n'
                
                # Fila de Criterio de evaluación
                html_tabla += '<tr>\n'
                html_tabla += '<td style="background-color: #1E1E1E; color: #E0E0E0; padding: 6px; border: 1px solid #555; font-weight: bold;">Criterio de evaluación:</td>\n'
                html_tabla += f'<td style="background-color: #1E1E1E; color: #E0E0E0; padding: 6px; border: 1px solid #555;">{escapar_html(criterio)}</td>\n'
                html_tabla += '</tr>\n'
                
                # Fila de Principales actividades
                html_tabla += '<tr>\n'
                html_tabla += '<td style="background-color: #1E1E1E; color: #E0E0E0; padding: 6px; border: 1px solid #555; font-weight: bold;">Principales actividades de aprendizaje:</td>\n'
                html_tabla += f'<td style="background-color: #1E1E1E; color: #E0E0E0; padding: 6px; border: 1px solid #555;">{escapar_html(actividades)}</td>\n'
                html_tabla += '</tr>\n'
                
                html_tabla += '</table>\n'
                html_tabla += '</td>\n'
        
        html_tabla += '</tr>\n'
    
    html_tabla += '</table>\n'
    return html_tabla

# Función para procesar y formatear el contenido de unidad didáctica
def formatear_unidad_didactica(contenido_raw, area_curricular, num_sesiones=4):
    """
    Procesa el contenido generado y lo estructura como una unidad didáctica profesional
    
    Args:
        contenido_raw: Contenido generado por la IA
        area_curricular: Área curricular
        num_sesiones: Número de sesiones para la tabla dinámica (por defecto 4)
    """
    # Normalizar las tablas antes de formatear
    contenido_normalizado = normalizar_tabla_para_streamlit(contenido_raw)
    
    # Convertir tablas markdown a HTML con los mismos estilos
    contenido_con_tablas_html = convertir_tablas_markdown_a_html(contenido_normalizado)
    
    # Extraer datos de SECUENCIA DE SESIONES (Título, Criterio, Principales actividades) del contenido generado
    def extraer_secuencia_sesiones(contenido, num_sesiones_esperadas=4):
        """
        Extrae por cada sesión: titulo, criterio de evaluación, principales actividades.
        Retorna lista de dicts con keys titulo, criterio, actividades.
        """
        resultado = []
        lineas = contenido.split('\n')
        contenido_secuencia = []
        dentro_secuencia = False
        
        for linea in lineas:
            linea_stripped = linea.strip()
            if 'SECUENCIA' in linea_stripped.upper() and 'SESION' in linea_stripped.upper():
                dentro_secuencia = True
                if '|' in linea_stripped:
                    partes = linea_stripped.split('|')
                    if len(partes) >= 3:
                        contenido_secuencia.append(partes[2].strip())
                    elif len(partes) >= 2:
                        contenido_secuencia.append(partes[1].strip())
                continue
            if dentro_secuencia and '|' in linea_stripped:
                partes = linea_stripped.split('|')
                if len(partes) >= 3:
                    item = (partes[1] or "").strip()
                    if item and '**' in item and 'SECUENCIA' not in item.upper():
                        break
                    contenido_secuencia.append(partes[2].strip())
                elif len(partes) >= 2:
                    contenido_secuencia.append(partes[1].strip())
            elif dentro_secuencia and linea_stripped and not linea_stripped.startswith('|'):
                if contenido_secuencia:
                    contenido_secuencia[-1] += " " + linea_stripped
            elif dentro_secuencia and not linea_stripped and contenido_secuencia:
                break
        
        texto = " ".join(contenido_secuencia).replace("  ", " ")
        
        for n in range(1, num_sesiones_esperadas + 1):
            bloque = ""
            patron_inicio = re.compile(r"Sesi[oó]n\s+" + str(n) + r"\s*[:\-]\s*", re.IGNORECASE)
            patron_siguiente = re.compile(r"Sesi[oó]n\s+" + str(n + 1) + r"\s*[:\-]\s*", re.IGNORECASE)
            match = patron_inicio.search(texto)
            if match:
                inicio = match.end()
                match_sig = patron_siguiente.search(texto[inicio:])
                if match_sig:
                    bloque = texto[inicio:inicio + match_sig.start()].strip()
                else:
                    bloque = texto[inicio:].strip()
            else:
                resultado.append({"titulo": "", "criterio": "", "actividades": ""})
                continue
            
            titulo = ""
            criterio = ""
            actividades = ""
            # Bloque tiene forma: "Título: X. Criterio de evaluación: Y. Principales actividades: Z."
            resto = bloque
            for sep in ("Título:", "Titulo:"):
                if sep in resto:
                    resto = resto.split(sep, 1)[1].strip()
                    for sep_c in ("Criterio de evaluación:", "Criterio de evaluacion:"):
                        if sep_c in resto:
                            titulo = resto.split(sep_c, 1)[0].strip()
                            resto = resto.split(sep_c, 1)[1].strip()
                            if "Principales actividades" in resto:
                                partes_act = re.split(r"Principales actividades\s*:?\s*", resto, 1, flags=re.IGNORECASE)
                                criterio = (partes_act[0] or "").strip()
                                actividades = (partes_act[-1] or "").strip() if len(partes_act) > 1 else ""
                            else:
                                criterio = resto.strip()
                            break
                    else:
                        titulo = resto[:500].strip()
                    break
            if not titulo and bloque:
                titulo = bloque[:400].strip()
            resultado.append({"titulo": (titulo[:500] if titulo else ""), "criterio": (criterio[:600] if criterio else ""), "actividades": (actividades[:800] if actividades else "")})
        
        while len(resultado) < num_sesiones_esperadas:
            resultado.append({"titulo": "", "criterio": "", "actividades": ""})
        return resultado[:num_sesiones_esperadas]
    
    sesiones_data = extraer_secuencia_sesiones(contenido_raw, num_sesiones)
    if not any(s.get("titulo") or s.get("criterio") or s.get("actividades") for s in sesiones_data):
        sesiones_data = extraer_secuencia_sesiones(contenido_normalizado, num_sesiones)
    
    # Generar tabla dinámica de secuencia de sesiones con datos extraídos
    tabla_secuencia_sesiones = generar_tabla_secuencia_sesiones(num_sesiones, sesiones_data=sesiones_data)
    
    # Extraer datos de la tabla de competencias transversales del contenido generado
    def extraer_competencias_transversales(contenido):
        """
        Extrae los datos de la tabla de competencias transversales del contenido generado.
        Retorna un diccionario con los estándares e instrumentos para cada competencia.
        """
        competencias_data = {
            "Se desenvuelve en los entornos virtuales generados por las TIC.": {
                "estandar": "",
                "instrumento": ""
            },
            "Gestiona su aprendizaje de manera autónoma.": {
                "estandar": "",
                "instrumento": ""
            }
        }
        
        # Buscar la tabla de competencias transversales en el contenido
        lineas = contenido.split('\n')
        dentro_tabla_comp_trans = False
        
        for i, linea in enumerate(lineas):
            linea_stripped = linea.strip()
            
            # Detectar inicio de tabla de competencias transversales (puede estar en encabezado o en contenido)
            if '|' in linea_stripped and ('COMPETENCIAS TRANSVERSALES' in linea_stripped.upper() or 
                                         ('ESTÁNDARES' in linea_stripped.upper() and 'INSTRUMENTO' in linea_stripped.upper())):
                dentro_tabla_comp_trans = True
                # Verificar si esta línea ya tiene datos (no es solo el encabezado)
                partes = linea_stripped.split('|')
                fila = [celda.strip() for celda in partes]
                # Eliminar celdas vacías al inicio y final
                while fila and not fila[0]:
                    fila.pop(0)
                while fila and not fila[-1]:
                    fila.pop()
                
                # Si tiene 3 columnas y contiene una competencia, procesarla
                if len(fila) == 3 and any(palabra in fila[0].upper() for palabra in ['SE DESENVUELVE', 'GESTIONA']):
                    competencia = fila[0]
                    estandar = fila[1] if len(fila) > 1 else ""
                    instrumento = fila[2] if len(fila) > 2 else ""
                    
                    if "Se desenvuelve" in competencia or "entornos virtuales" in competencia:
                        competencias_data["Se desenvuelve en los entornos virtuales generados por las TIC."]["estandar"] = estandar
                        competencias_data["Se desenvuelve en los entornos virtuales generados por las TIC."]["instrumento"] = instrumento
                    elif "Gestiona" in competencia or "aprendizaje" in competencia and "autónoma" in competencia:
                        competencias_data["Gestiona su aprendizaje de manera autónoma."]["estandar"] = estandar
                        competencias_data["Gestiona su aprendizaje de manera autónoma."]["instrumento"] = instrumento
                continue
            
            if dentro_tabla_comp_trans:
                # Detectar si es una fila de datos (tiene 3 columnas y contiene las competencias)
                if '|' in linea_stripped and linea_stripped.count('|') >= 3:
                    partes = linea_stripped.split('|')
                    fila = [celda.strip() for celda in partes]
                    # Eliminar celdas vacías al inicio y final
                    while fila and not fila[0]:
                        fila.pop(0)
                    while fila and not fila[-1]:
                        fila.pop()
                    
                    # Verificar si es una fila de datos (3 columnas y contiene una de las competencias)
                    if len(fila) == 3:
                        competencia = fila[0]
                        estandar = fila[1] if len(fila) > 1 else ""
                        instrumento = fila[2] if len(fila) > 2 else ""
                        
                        # Verificar si corresponde a una de las competencias conocidas
                        if "Se desenvuelve" in competencia or ("entornos virtuales" in competencia and "TIC" in competencia):
                            if not competencias_data["Se desenvuelve en los entornos virtuales generados por las TIC."]["estandar"]:
                                competencias_data["Se desenvuelve en los entornos virtuales generados por las TIC."]["estandar"] = estandar
                            if not competencias_data["Se desenvuelve en los entornos virtuales generados por las TIC."]["instrumento"]:
                                competencias_data["Se desenvuelve en los entornos virtuales generados por las TIC."]["instrumento"] = instrumento
                        elif "Gestiona" in competencia and ("aprendizaje" in competencia or "autónoma" in competencia):
                            if not competencias_data["Gestiona su aprendizaje de manera autónoma."]["estandar"]:
                                competencias_data["Gestiona su aprendizaje de manera autónoma."]["estandar"] = estandar
                            if not competencias_data["Gestiona su aprendizaje de manera autónoma."]["instrumento"]:
                                competencias_data["Gestiona su aprendizaje de manera autónoma."]["instrumento"] = instrumento
                elif linea_stripped and '|' not in linea_stripped:
                    # Línea fuera de tabla, puede ser el fin
                    if not any(palabra in linea_stripped.upper() for palabra in ['COMPETENCIAS', 'TRANSVERSALES', 'ESTÁNDARES', 'INSTRUMENTO', 'SE DESENVUELVE', 'GESTIONA']):
                        # Ya no estamos en la tabla
                        break
        
        return competencias_data
    
    # Extraer datos de competencias transversales
    competencias_trans_data = extraer_competencias_transversales(contenido_raw)
    
    # Extraer datos de PROPÓSITOS DE APRENDIZAJE para la tabla de 5 columnas
    def extraer_propositos_aprendizaje(contenido):
        """
        Extrae los datos de PROPÓSITOS DE APRENDIZAJE del contenido generado.
        Retorna un diccionario con competencias, capacidades, criterios, evidencias e instrumentos.
        """
        propositos_data = {
            "competencias": "",
            "capacidades": "",
            "criterios": "",
            "evidencias": "",
            "instrumentos": ""
        }
        
        lineas = contenido.split('\n')
        dentro_propositos = False
        contenido_propositos = []
        
        for i, linea in enumerate(lineas):
            linea_stripped = linea.strip()
            
            # Detectar inicio de PROPÓSITOS DE APRENDIZAJE
            if 'PROPÓSITOS DE APRENDIZAJE' in linea_stripped.upper() or 'III. PROPÓSITOS' in linea_stripped.upper():
                dentro_propositos = True
                # Extraer contenido de esta línea si está presente
                if '|' in linea_stripped:
                    partes = linea_stripped.split('|')
                    if len(partes) >= 3:
                        contenido_celda = partes[2].strip() if len(partes) > 2 else ""
                        if contenido_celda:
                            contenido_propositos.append(contenido_celda)
                continue
            
            # Si estamos dentro de PROPÓSITOS DE APRENDIZAJE, seguir extrayendo
            if dentro_propositos:
                if '|' in linea_stripped:
                    partes = linea_stripped.split('|')
                    if len(partes) >= 3:
                        item_col = partes[1].strip()
                        contenido_celda = partes[2].strip() if len(partes) > 2 else ""
                        
                        # Verificar si encontramos otro ítem (fin de PROPÓSITOS DE APRENDIZAJE)
                        if item_col and '**' in item_col and 'PROPÓSITOS' not in item_col.upper():
                            break
                        
                        # Si hay contenido en la columna derecha, agregarlo
                        if contenido_celda:
                            contenido_propositos.append(contenido_celda)
                elif linea_stripped and not linea_stripped.startswith('|'):
                    # Línea fuera de tabla, puede ser contenido multilínea
                    if contenido_propositos:
                        contenido_propositos[-1] += " " + linea_stripped
                elif not linea_stripped and contenido_propositos:
                    # Línea vacía, puede indicar fin de sección
                    break
        
        # Procesar el contenido extraído para separar competencias, capacidades, criterios, etc.
        contenido_completo = " ".join(contenido_propositos)
        
        # Extraer Competencias
        if 'Competencias:' in contenido_completo or 'COMPETENCIAS:' in contenido_completo:
            partes = contenido_completo.split('Capacidades:')
            if len(partes) > 0:
                competencias_texto = partes[0].replace('Competencias:', '').replace('COMPETENCIAS:', '').strip()
                propositos_data["competencias"] = competencias_texto.split('Criterios')[0].split('CRITERIOS')[0].strip()
        
        # Extraer Capacidades
        if 'Capacidades:' in contenido_completo or 'CAPACIDADES:' in contenido_completo:
            partes = contenido_completo.split('Criterios de evaluación:')
            if len(partes) > 0:
                capacidades_texto = partes[0]
                if 'Capacidades:' in capacidades_texto:
                    capacidades_texto = capacidades_texto.split('Capacidades:')[1].strip()
                elif 'CAPACIDADES:' in capacidades_texto:
                    capacidades_texto = capacidades_texto.split('CAPACIDADES:')[1].strip()
                propositos_data["capacidades"] = capacidades_texto.split('Contenidos')[0].split('CONTENIDOS')[0].strip()
        
        # Extraer Criterios de evaluación
        if 'Criterios de evaluación:' in contenido_completo or 'CRITERIOS DE EVALUACIÓN:' in contenido_completo or 'Criterios:' in contenido_completo:
            partes = contenido_completo.split('Contenidos:')
            if len(partes) > 0:
                criterios_texto = partes[0]
                if 'Criterios de evaluación:' in criterios_texto:
                    criterios_texto = criterios_texto.split('Criterios de evaluación:')[1].strip()
                elif 'CRITERIOS DE EVALUACIÓN:' in criterios_texto:
                    criterios_texto = criterios_texto.split('CRITERIOS DE EVALUACIÓN:')[1].strip()
                elif 'Criterios:' in criterios_texto:
                    criterios_texto = criterios_texto.split('Criterios:')[1].strip()
                propositos_data["criterios"] = criterios_texto.split('Evidencia')[0].split('EVIDENCIA')[0].strip()
        
        # Extraer Evidencias de aprendizaje
        if 'Evidencia de aprendizaje:' in contenido_completo or 'EVIDENCIA DE APRENDIZAJE:' in contenido_completo or 'Evidencia:' in contenido_completo:
            partes = contenido_completo.split('Instrumento de evaluación:')
            if len(partes) > 0:
                evidencias_texto = partes[0]
                if 'Evidencia de aprendizaje:' in evidencias_texto:
                    evidencias_texto = evidencias_texto.split('Evidencia de aprendizaje:')[1].strip()
                elif 'EVIDENCIA DE APRENDIZAJE:' in evidencias_texto:
                    evidencias_texto = evidencias_texto.split('EVIDENCIA DE APRENDIZAJE:')[1].strip()
                elif 'Evidencia:' in evidencias_texto:
                    evidencias_texto = evidencias_texto.split('Evidencia:')[1].strip()
                propositos_data["evidencias"] = evidencias_texto.split('Instrumento')[0].strip()
        
        # Extraer Instrumentos de evaluación
        if 'Instrumento de evaluación:' in contenido_completo or 'INSTRUMENTO DE EVALUACIÓN:' in contenido_completo or 'Instrumento:' in contenido_completo:
            partes = contenido_completo.split('Instrumento de evaluación:')
            if len(partes) > 1:
                propositos_data["instrumentos"] = partes[1].strip()
            elif 'INSTRUMENTO DE EVALUACIÓN:' in contenido_completo:
                partes = contenido_completo.split('INSTRUMENTO DE EVALUACIÓN:')
                if len(partes) > 1:
                    propositos_data["instrumentos"] = partes[1].strip()
            elif 'Instrumento:' in contenido_completo:
                partes = contenido_completo.split('Instrumento:')
                if len(partes) > 1:
                    propositos_data["instrumentos"] = partes[1].strip()
        
        return propositos_data
    
    # Extraer datos de propósitos de aprendizaje
    propositos_data = extraer_propositos_aprendizaje(contenido_raw)
    
    # Si no se encontraron datos, intentar extraer del contenido normalizado también
    if not propositos_data.get("competencias") and not propositos_data.get("capacidades"):
        propositos_data = extraer_propositos_aprendizaje(contenido_normalizado)
    
    # Extraer datos de VALORES Y ENFOQUES TRANSVERSALES para la tabla de 4 columnas
    def extraer_valores_enfoques_transversales(contenido):
        """
        Extrae el contenido de VALORES Y ENFOQUES TRANSVERSALES del contenido generado.
        Retorna dict con valores, enfoques_transversales, comportamientos_observables, instrumento.
        """
        data = {
            "valores": "",
            "enfoques_transversales": "",
            "comportamientos_observables": "",
            "instrumento": ""
        }
        
        lineas = contenido.split('\n')
        dentro_valores = False
        contenido_valores = []
        
        for linea in lineas:
            linea_stripped = linea.strip()
            
            if 'VALORES Y ENFOQUES TRANSVERSALES' in linea_stripped.upper():
                dentro_valores = True
                if '|' in linea_stripped:
                    partes = linea_stripped.split('|')
                    if len(partes) >= 3:
                        contenido_celda = partes[2].strip() if len(partes) > 2 else ""
                        if contenido_celda:
                            contenido_valores.append(contenido_celda)
                continue
            
            if dentro_valores:
                if '|' in linea_stripped:
                    partes = linea_stripped.split('|')
                    if len(partes) >= 3:
                        item_col = partes[1].strip()
                        contenido_celda = partes[2].strip() if len(partes) > 2 else ""
                        if item_col and '**' in item_col and 'VALORES' not in item_col.upper() and 'ENFOQUES' not in item_col.upper():
                            break
                        if contenido_celda:
                            contenido_valores.append(contenido_celda)
                elif linea_stripped and not linea_stripped.startswith('|'):
                    if contenido_valores:
                        contenido_valores[-1] += " " + linea_stripped
                elif not linea_stripped and contenido_valores:
                    break
        
        texto_completo = " ".join(contenido_valores).replace("  ", " ").strip()
        if not texto_completo:
            return data
        
        # Extraer sección Valores (después de "Valores:" hasta "Enfoques:" o "Con comportamientos")
        for sep in ('Valores:', 'VALORES:'):
            if sep in texto_completo:
                idx = texto_completo.find(sep)
                resto = texto_completo[idx + len(sep):].strip()
                for corte in ('Enfoques:', 'ENFOQUES:', 'Con comportamientos', 'con comportamientos'):
                    if corte in resto:
                        resto = resto.split(corte)[0].strip()
                # Quitar instrucción del prompt y quedarnos con el listado de valores
                for instr in ('DEBES usar SOLO estos 13 valores (sin agregar ni omitir):', 'DEBES usar SOLO estos 13 valores (sin agregar ni omitir):', 'usar SOLO estos 13 valores:', 'estos 13 valores (sin agregar ni omitir):'):
                    if instr in resto:
                        resto = resto.split(instr)[-1].strip()
                if resto.startswith(':'):
                    resto = resto[1:].strip()
                data["valores"] = resto[:700].strip() if len(resto) > 700 else resto
                break
        
        # Extraer sección Enfoques transversales (después de "Enfoques:" hasta "Con comportamientos" o punto)
        for sep in ('Enfoques:', 'ENFOQUES:'):
            if sep in texto_completo:
                idx = texto_completo.find(sep)
                resto = texto_completo[idx + len(sep):].strip()
                for corte in ('Con comportamientos', 'con comportamientos', 'Pueden incluir'):
                    if corte in resto:
                        resto = resto.split(corte)[0].strip()
                data["enfoques_transversales"] = resto[:600].strip() if len(resto) > 600 else resto
                break
        
        # Comportamientos observables: texto que menciona comportamientos o lo que sigue a "comportamientos observables"
        if 'comportamientos observables' in texto_completo.lower():
            idx = texto_completo.lower().find('comportamientos observables')
            # Tomar desde "para cada uno" o desde "observables" si hay más texto después
            candidato = texto_completo[idx:].strip()
            if candidato.startswith('comportamientos observables para cada uno'):
                candidato = candidato[len('comportamientos observables para cada uno'):].strip()
            elif candidato.startswith('comportamientos observables.'):
                candidato = candidato[len('comportamientos observables.'):].strip()
            else:
                candidato = candidato[len('comportamientos observables'):].lstrip(' :.').strip()
            if candidato:
                data["comportamientos_observables"] = candidato[:800] + ("..." if len(candidato) > 800 else "")
        
        # Si no se extrajo nada, repartir el texto: primera mitad enfoques, segunda comportamientos
        if not data["enfoques_transversales"] and not data["comportamientos_observables"]:
            mitad = len(texto_completo) // 2
            data["enfoques_transversales"] = texto_completo[:mitad].strip()[:600]
            data["comportamientos_observables"] = (texto_completo[mitad:].strip()[:800] or "")
        elif not data["comportamientos_observables"] and data["enfoques_transversales"]:
            data["comportamientos_observables"] = texto_completo[:800] if len(texto_completo) > 800 else texto_completo
        
        # Instrumento por defecto si no se indica en el contenido
        data["instrumento"] = data.get("instrumento") or "Lista de cotejo o rúbrica de actitudes para valores y enfoques transversales"
        
        return data
    
    valores_enfoques_data = extraer_valores_enfoques_transversales(contenido_raw)
    if not valores_enfoques_data.get("valores") and not valores_enfoques_data.get("enfoques_transversales") and not valores_enfoques_data.get("comportamientos_observables"):
        valores_enfoques_data = extraer_valores_enfoques_transversales(contenido_normalizado)
    
    # Cuadro de DATOS INFORMATIVOS en formato markdown para Streamlit
    # Estructura: 4 columnas en las primeras 3 filas, luego filas con rowspan y colspan
    # Colores adaptados para modo oscuro
    cuadro_datos = f"""
### I. DATOS INFORMATIVOS

<table style="width: 100%; border-collapse: collapse;">
<tr>
<td style="background-color: #2C3E50; color: #FFFFFF; padding: 8px; border: 1px solid #555; font-weight: bold;">DRE / UGEL</td>
<td style="background-color: #1E1E1E; color: #E0E0E0; padding: 8px; border: 1px solid #555;"></td>
<td style="background-color: #2C3E50; color: #FFFFFF; padding: 8px; border: 1px solid #555; font-weight: bold;">Institución educativa</td>
<td style="background-color: #1E1E1E; color: #E0E0E0; padding: 8px; border: 1px solid #555;"></td>
</tr>
<tr>
<td style="background-color: #2C3E50; color: #FFFFFF; padding: 8px; border: 1px solid #555; font-weight: bold;">Área curricular</td>
<td style="background-color: #1E1E1E; color: #E0E0E0; padding: 8px; border: 1px solid #555;"></td>
<td style="background-color: #2C3E50; color: #FFFFFF; padding: 8px; border: 1px solid #555; font-weight: bold;">Grado y secciones</td>
<td style="background-color: #1E1E1E; color: #E0E0E0; padding: 8px; border: 1px solid #555;"></td>
</tr>
<tr>
<td style="background-color: #2C3E50; color: #FFFFFF; padding: 8px; border: 1px solid #555; font-weight: bold;">Fecha de inicio</td>
<td style="background-color: #1E1E1E; color: #E0E0E0; padding: 8px; border: 1px solid #555;"></td>
<td style="background-color: #2C3E50; color: #FFFFFF; padding: 8px; border: 1px solid #555; font-weight: bold;">Fecha de término</td>
<td style="background-color: #1E1E1E; color: #E0E0E0; padding: 8px; border: 1px solid #555;"></td>
</tr>
<tr>
<td rowspan="2" style="background-color: #2C3E50; color: #FFFFFF; padding: 8px; border: 1px solid #555; font-weight: bold; vertical-align: top;">Coordinación pedagógica</td>
<td colspan="3" style="background-color: #1E1E1E; color: #E0E0E0; padding: 8px; border: 1px solid #555;"></td>
</tr>
<tr>
<td colspan="3" style="background-color: #1E1E1E; color: #E0E0E0; padding: 8px; border: 1px solid #555;"></td>
</tr>
<tr>
<td style="background-color: #2C3E50; color: #FFFFFF; padding: 8px; border: 1px solid #555; font-weight: bold;">Docente(s)</td>
<td colspan="3" style="background-color: #1E1E1E; color: #E0E0E0; padding: 8px; border: 1px solid #555;"></td>
</tr>
</table>

---
"""
    
    contenido_formateado = f"""
# 📚 UNIDAD DIDÁCTICA

## 📋 ÁREA CURRICULAR: {area_curricular}

---

{cuadro_datos}

### II. SITUACIÓN SIGNIFICATIVA

<div style="color: #E0E0E0; padding: 10px 0;">
[El contenido de la situación significativa se generará aquí]
</div>

---

### III. PROPÓSITOS DE APRENDIZAJE

<div style="color: #E0E0E0; padding: 10px 0;">
[El contenido de los propósitos de aprendizaje se generará aquí]
</div>

---

### COMPETENCIAS TRANSVERSALES

<table style="width: 100%; border-collapse: collapse;">
<tr>
<td style="background-color: #2C3E50; color: #FFFFFF; padding: 8px; border: 1px solid #555; font-weight: bold; text-align: center;">Competencias transversales</td>
<td style="background-color: #2C3E50; color: #FFFFFF; padding: 8px; border: 1px solid #555; font-weight: bold; text-align: center;">Estándares de aprendizaje</td>
<td style="background-color: #2C3E50; color: #FFFFFF; padding: 8px; border: 1px solid #555; font-weight: bold; text-align: center;">Instrumento</td>
</tr>
<tr>
<td style="background-color: #1E1E1E; color: #E0E0E0; padding: 8px; border: 1px solid #555;">Se desenvuelve en los entornos virtuales generados por las TIC.</td>
<td style="background-color: #1E1E1E; color: #E0E0E0; padding: 8px; border: 1px solid #555;">{competencias_trans_data.get("Se desenvuelve en los entornos virtuales generados por las TIC.", {}).get("estandar", "")}</td>
<td style="background-color: #1E1E1E; color: #E0E0E0; padding: 8px; border: 1px solid #555;">{competencias_trans_data.get("Se desenvuelve en los entornos virtuales generados por las TIC.", {}).get("instrumento", "")}</td>
</tr>
<tr>
<td style="background-color: #1E1E1E; color: #E0E0E0; padding: 8px; border: 1px solid #555;">Gestiona su aprendizaje de manera autónoma.</td>
<td style="background-color: #1E1E1E; color: #E0E0E0; padding: 8px; border: 1px solid #555;">{competencias_trans_data.get("Gestiona su aprendizaje de manera autónoma.", {}).get("estandar", "")}</td>
<td style="background-color: #1E1E1E; color: #E0E0E0; padding: 8px; border: 1px solid #555;">{competencias_trans_data.get("Gestiona su aprendizaje de manera autónoma.", {}).get("instrumento", "")}</td>
</tr>
</table>

---

<table style="width: 100%; border-collapse: collapse;">
<tr>
<td style="background-color: #2C3E50; color: #FFFFFF; padding: 8px; border: 1px solid #555; font-weight: bold; text-align: center;">Competencias de área</td>
<td style="background-color: #2C3E50; color: #FFFFFF; padding: 8px; border: 1px solid #555; font-weight: bold; text-align: center;">Capacidades</td>
<td style="background-color: #2C3E50; color: #FFFFFF; padding: 8px; border: 1px solid #555; font-weight: bold; text-align: center;">Criterios de evaluación</td>
<td style="background-color: #2C3E50; color: #FFFFFF; padding: 8px; border: 1px solid #555; font-weight: bold; text-align: center;">Evidencia de aprendizaje</td>
<td style="background-color: #2C3E50; color: #FFFFFF; padding: 8px; border: 1px solid #555; font-weight: bold; text-align: center;">Instrumento</td>
</tr>
<tr>
<td style="background-color: #1E1E1E; color: #E0E0E0; padding: 8px; border: 1px solid #555;">{propositos_data.get("competencias", "")}</td>
<td style="background-color: #1E1E1E; color: #E0E0E0; padding: 8px; border: 1px solid #555;">{propositos_data.get("capacidades", "")}</td>
<td style="background-color: #1E1E1E; color: #FF6B6B; padding: 8px; border: 1px solid #555;">{propositos_data.get("criterios", "Hay que recordar que los criterios se desprenden de los estándares de aprendizaje.")}</td>
<td style="background-color: #1E1E1E; color: #E0E0E0; padding: 8px; border: 1px solid #555;">{propositos_data.get("evidencias", "")}</td>
<td style="background-color: #1E1E1E; color: #E0E0E0; padding: 8px; border: 1px solid #555;">{propositos_data.get("instrumentos", "")}</td>
</tr>
</table>

---

<table style="width: 100%; border-collapse: collapse;">
<tr>
<td style="background-color: #2C3E50; color: #FFFFFF; padding: 8px; border: 1px solid #555; font-weight: bold; text-align: center;">Valor</td>
<td style="background-color: #2C3E50; color: #FFFFFF; padding: 8px; border: 1px solid #555; font-weight: bold; text-align: center;">Enfoques transversales</td>
<td style="background-color: #2C3E50; color: #FFFFFF; padding: 8px; border: 1px solid #555; font-weight: bold; text-align: center;">Comportamientos observables</td>
<td style="background-color: #2C3E50; color: #FFFFFF; padding: 8px; border: 1px solid #555; font-weight: bold; text-align: center;">Instrumento</td>
</tr>
<tr>
<td rowspan="2" style="background-color: #1E1E1E; color: #E0E0E0; padding: 8px; border: 1px solid #555; vertical-align: top;">{valores_enfoques_data.get("valores", "")}</td>
<td style="background-color: #1E1E1E; color: #E0E0E0; padding: 8px; border: 1px solid #555;">{valores_enfoques_data.get("enfoques_transversales", "")}</td>
<td style="background-color: #1E1E1E; color: #E0E0E0; padding: 8px; border: 1px solid #555;">{valores_enfoques_data.get("comportamientos_observables", "")}</td>
<td style="background-color: #1E1E1E; color: #E0E0E0; padding: 8px; border: 1px solid #555;">{valores_enfoques_data.get("instrumento", "")}</td>
</tr>
<tr>
<td style="background-color: #1E1E1E; color: #E0E0E0; padding: 8px; border: 1px solid #555;"></td>
<td style="background-color: #1E1E1E; color: #E0E0E0; padding: 8px; border: 1px solid #555;"></td>
<td style="background-color: #1E1E1E; color: #E0E0E0; padding: 8px; border: 1px solid #555;"></td>
</tr>
</table>

---

{tabla_secuencia_sesiones}

---

### 📝 NOTAS METODOLÓGICAS

Esta unidad didáctica ha sido diseñada siguiendo los lineamientos del Currículo Nacional de la Educación Básica del Perú.

**Recomendaciones de implementación:**
- Considerar el contexto sociocultural de los estudiantes
- Adaptar las estrategias según los ritmos de aprendizaje
- Integrar recursos tecnológicos disponibles
- Promover el aprendizaje colaborativo

---

*Documento generado automáticamente por el Sistema de IA Educativa*
"""
    # Títulos de sesiones (los mismos que aparecen en IV. SECUENCIA DE SESIONES) para el selector en Sesión de Aprendizaje
    titulos_sesiones = [s.get("titulo", "").strip() for s in sesiones_data if s.get("titulo") and str(s.get("titulo", "")).strip()]
    return contenido_formateado, titulos_sesiones

# Función para procesar y formatear el contenido de sesión de aprendizaje
def formatear_sesion_aprendizaje(contenido_raw, titulo_unidad, titulo_sesion, nivel, grado, seccion, duracion, area_curricular=""):
    """
    Procesa el contenido generado y lo estructura como una sesión de aprendizaje profesional.
    Incluye la tabla I. DATOS INFORMATIVOS rellenada con los datos del formulario y lo generado.
    """
    # Normalizar las tablas antes de formatear
    contenido_normalizado = normalizar_tabla_para_streamlit(contenido_raw)
    
    # Extraer Situación significativa del contenido generado (tabla | **SITUACIÓN SIGNIFICATIVA** | contenido |)
    def extraer_situacion_significativa(contenido):
        lineas = contenido.split('\n')
        for linea in lineas:
            linea_stripped = linea.strip()
            if '|' in linea_stripped and 'SITUACIÓN SIGNIFICATIVA' in linea_stripped.upper():
                partes = [p.strip() for p in linea_stripped.split('|')]
                while partes and not partes[0]:
                    partes.pop(0)
                while partes and not partes[-1]:
                    partes.pop()
                if len(partes) >= 2:
                    return partes[1].replace('**', '').strip()
                if len(partes) >= 3:
                    return partes[2].replace('**', '').strip()
        return ""
    
    situacion_significativa = extraer_situacion_significativa(contenido_raw)
    if not situacion_significativa:
        situacion_significativa = extraer_situacion_significativa(contenido_normalizado)
    
    # Extraer PROPÓSITOS DE APRENDIZAJE (Competencias, Capacidades, Criterios, Contenidos, Evidencia, Instrumento)
    def extraer_propositos_sesion(contenido):
        data = {"competencias": "", "capacidades": "", "criterios": "", "contenidos": "", "evidencia": "", "instrumento": ""}
        for linea in contenido.split('\n'):
            l = linea.strip()
            if '|' in l and 'PROPÓSITOS DE APRENDIZAJE' in l.upper():
                partes = [p.strip() for p in l.split('|')]
                while partes and not partes[0]:
                    partes.pop(0)
                while partes and not partes[-1]:
                    partes.pop()
                if len(partes) < 2:
                    continue
                texto = partes[1].replace('**', '')
                for clave, sep in [
                    ("competencias", "Competencias:"),
                    ("capacidades", "Capacidades:"),
                    ("criterios", "Criterios de evaluación:"),
                    ("contenidos", "Contenidos:"),
                    ("evidencia", "Evidencia de aprendizaje:"),
                    ("instrumento", "Instrumento de evaluación:")
                ]:
                    if sep in texto:
                        idx = texto.find(sep)
                        resto = texto[idx + len(sep):].strip()
                        siguiente = None
                        for s in ["Competencias:", "Capacidades:", "Criterios de evaluación:", "Contenidos:", "Evidencia de aprendizaje:", "Instrumento de evaluación:"]:
                            if s != sep and s in resto:
                                siguiente = resto.find(s)
                                break
                        data[clave] = (resto[:siguiente].strip() if siguiente is not None else resto)[:800]
                break
        return data
    
    # Extraer COMPETENCIAS TRANSVERSALES (para las dos competencias fijas: capacidad transversal + desempeño transversal)
    def extraer_comp_transversales_sesion(contenido):
        comp1 = "Se desenvuelve en los entornos virtuales generados por las TIC."
        comp2 = "Gestiona su aprendizaje de manera autónoma."
        resultado = [
            {"competencia": comp1, "capacidad": "", "desempeno": ""},
            {"competencia": comp2, "capacidad": "", "desempeno": ""}
        ]
        texto_celda = ""
        for linea in contenido.split('\n'):
            l = linea.strip()
            if '|' in l and 'COMPETENCIAS TRANSVERSALES' in l.upper():
                partes = [p.strip() for p in l.split('|')]
                while partes and not partes[0]:
                    partes.pop(0)
                while partes and not partes[-1]:
                    partes.pop()
                if len(partes) >= 2:
                    texto_celda = partes[1].replace('**', '')
                break
        if not texto_celda:
            return resultado
        # Puede haber dos bloques (uno por competencia) separados por la segunda competencia
        bloques = re.split(re.escape(comp2), texto_celda, 1)
        for i, bloque in enumerate([bloques[0]] if len(bloques) == 1 else bloques):
            if i == 1 and len(bloques) > 1:
                bloque = comp2 + bloque
            if "Capacidad transversal:" in bloque:
                resto = bloque.split("Capacidad transversal:", 1)[1]
                if "Desempeño transversal:" in resto:
                    resultado[i]["capacidad"] = resto.split("Desempeño transversal:")[0].strip()[:400]
                    resultado[i]["desempeno"] = resto.split("Desempeño transversal:")[1].strip()[:500]
                else:
                    resultado[i]["capacidad"] = resto.strip()[:400]
            if i == 0 and len(bloques) == 1:
                resultado[1]["capacidad"] = resultado[0]["capacidad"]
                resultado[1]["desempeno"] = resultado[0]["desempeno"]
                break
        return resultado
    
    # Extraer ENFOQUE TRANSVERSAL (Valor priorizado, Valor operativo, Comportamientos observables)
    def extraer_enfoque_transversal_sesion(contenido):
        data = {"valor_priorizado": "", "valor_operativo": "", "comportamientos": ""}
        for linea in contenido.split('\n'):
            l = linea.strip()
            if '|' in l and 'ENFOQUE TRANSVERSAL' in l.upper():
                partes = [p.strip() for p in l.split('|')]
                while partes and not partes[0]:
                    partes.pop(0)
                while partes and not partes[-1]:
                    partes.pop()
                if len(partes) < 2:
                    continue
                texto = partes[1].replace('**', '')
                if "Valor priorizado:" in texto:
                    data["valor_priorizado"] = texto.split("Valor priorizado:")[1].split("Valor operativo:")[0].strip()[:300]
                if "Valor operativo:" in texto:
                    data["valor_operativo"] = texto.split("Valor operativo:")[1].split("Comportamientos observables:")[0].strip()[:300]
                if "Comportamientos observables:" in texto:
                    data["comportamientos"] = texto.split("Comportamientos observables:")[1].strip()[:600]
                break
        return data
    
    propositos_sesion = extraer_propositos_sesion(contenido_raw)
    if not any(propositos_sesion.values()):
        propositos_sesion = extraer_propositos_sesion(contenido_normalizado)
    comp_trans_sesion = extraer_comp_transversales_sesion(contenido_raw)
    if not comp_trans_sesion[0].get("capacidad") and not comp_trans_sesion[0].get("desempeno"):
        comp_trans_sesion = extraer_comp_transversales_sesion(contenido_normalizado)
    enfoque_sesion = extraer_enfoque_transversal_sesion(contenido_raw)
    if not any(enfoque_sesion.values()):
        enfoque_sesion = extraer_enfoque_transversal_sesion(contenido_normalizado)
    
    # Extraer SECUENCIA DIDÁCTICA (Inicio, Desarrollo, Cierre) para tabla III
    def extraer_secuencia_didactica(contenido):
        data = {"inicio": "", "desarrollo": "", "cierre": "", "recursos": ""}
        lineas = contenido.split('\n')
        contenido_secuencia = []
        contenido_materiales = []
        dentro_secuencia = False
        dentro_materiales = False
        for i, linea in enumerate(lineas):
            l = linea.strip()
            if '|' in l and 'SECUENCIA DIDÁCTICA' in l.upper():
                dentro_secuencia = True
                partes = [p.strip() for p in l.split('|')]
                while partes and not partes[0]: partes.pop(0)
                while partes and not partes[-1]: partes.pop()
                # Tabla | ITEM | CONTENIDO | -> contenido en partes[2] si hay 3+ columnas, si no en partes[1]
                celda = (partes[2] if len(partes) >= 3 else partes[1] if len(partes) >= 2 else "").replace('**', '')
                if celda:
                    contenido_secuencia.append(celda)
                continue
            if dentro_secuencia:
                if '|' in l and '**' in l:
                    item = (l.split('|')[1] or "").strip() if len(l.split('|')) > 1 else ""
                    if item and 'SECUENCIA' not in item.upper():
                        dentro_secuencia = False
                        continue
                if '|' in l:
                    partes = [p.strip() for p in l.split('|')]
                    # Contenido de celda derecha: última parte con contenido o partes[2]
                    if len(partes) >= 3 and partes[2]:
                        contenido_secuencia.append(partes[2].replace('**', ''))
                    elif len(partes) >= 2 and partes[1] and (not partes[1].upper().startswith('SECUENCIA') or len(partes) == 2):
                        contenido_secuencia.append(partes[1].replace('**', ''))
                elif l and not l.startswith('|'):
                    if contenido_secuencia:
                        contenido_secuencia[-1] += "\n" + l
                    else:
                        contenido_secuencia.append(l)
            if '|' in l and 'MATERIALES Y RECURSOS' in l.upper():
                dentro_materiales = True
                partes = [p.strip() for p in l.split('|')]
                while partes and not partes[0]: partes.pop(0)
                while partes and not partes[-1]: partes.pop()
                if len(partes) >= 2:
                    contenido_materiales.append(partes[1].replace('**', ''))
                continue
            if dentro_materiales and '|' in l and '**' in l:
                item = (l.split('|')[1] or "").strip() if len(l.split('|')) > 1 else ""
                if item and 'MATERIALES' not in item.upper():
                    dentro_materiales = False
                elif len(l.split('|')) >= 2:
                    contenido_materiales.append(l.split('|')[1].strip().replace('**', ''))
        texto_sec = " ".join(contenido_secuencia).replace("\n", " ")
        if not texto_sec.strip():
            return data
        # Dividir por secciones Desarrollo y Cierre (palabras completas)
        m_d = re.search(r'\bDesarrollo\b', texto_sec, re.IGNORECASE)
        if m_d:
            idx_d = m_d.start()
            data["inicio"] = texto_sec[:idx_d].strip()
            resto = texto_sec[idx_d:]
            m_c = re.search(r'\bCierre\b', resto, re.IGNORECASE)
            if m_c:
                idx_c = m_c.start()
                data["desarrollo"] = resto[:idx_c].strip()
                data["cierre"] = resto[idx_c:].strip()
            else:
                data["desarrollo"] = resto.strip()
        else:
            data["inicio"] = texto_sec.strip()
        data["recursos"] = " ".join(contenido_materiales).replace("\n", " ").strip() if contenido_materiales else ""
        return data
    
    secuencia_didactica = extraer_secuencia_didactica(contenido_raw)
    if not secuencia_didactica.get("inicio") and not secuencia_didactica.get("desarrollo"):
        secuencia_didactica = extraer_secuencia_didactica(contenido_normalizado)
    
    # Extraer REFLEXIÓN SOBRE LA ACTIVIDAD (Dificultades, Mejoras, Ajustes) para tabla IV - todo el texto en una celda
    def extraer_reflexion_actividad(contenido):
        lineas = contenido.split('\n')
        contenido_reflex = []
        dentro = False
        for linea in lineas:
            l = linea.strip()
            if '|' in l and 'REFLEXIÓN' in l.upper() and 'ACTIVIDAD' in l.upper():
                dentro = True
                partes = [p.strip() for p in l.split('|')]
                while partes and not partes[0]: partes.pop(0)
                while partes and not partes[-1]: partes.pop()
                celda = (partes[2] if len(partes) >= 3 else partes[1] if len(partes) >= 2 else "").replace('**', '')
                if celda:
                    contenido_reflex.append(celda)
                continue
            if dentro:
                if '|' in l:
                    partes = [p.strip() for p in l.split('|')]
                    # Primera columna (índice 1 en markdown | col1 | col2 |)
                    primera_col = partes[1] if len(partes) > 1 else ""
                    # Solo salir si la primera columna tiene un nuevo ítem (ej. **MATERIALES**) que no es REFLEXIÓN
                    if primera_col and '**' in primera_col and 'REFLEXIÓN' not in primera_col.upper():
                        break
                    # Contenido: segunda columna (índice 2) o primera si solo hay una celda con contenido
                    contenido_celda = (partes[2] if len(partes) > 2 and partes[2] else partes[1] if len(partes) > 1 else "").replace('**', '')
                    if contenido_celda:
                        contenido_reflex.append(contenido_celda)
                elif l and not l.startswith('|'):
                    if contenido_reflex:
                        contenido_reflex[-1] += "\n" + l
                    else:
                        contenido_reflex.append(l)
        # Unir preservando saltos de línea para que se vean Dificultades, Mejoras, Ajustes
        if not contenido_reflex:
            return ""
        return "\n\n".join(contenido_reflex).strip()
    
    reflexion_actividad = extraer_reflexion_actividad(contenido_raw)
    if not reflexion_actividad:
        reflexion_actividad = extraer_reflexion_actividad(contenido_normalizado)
    # Fallback: buscar Inicio/Desarrollo/Cierre en todo el texto si la tabla no trajo contenido
    if not secuencia_didactica.get("inicio") and not secuencia_didactica.get("desarrollo") and not secuencia_didactica.get("cierre"):
        texto_completo = contenido_raw + "\n" + contenido_normalizado
        m_d = re.search(r'\bDesarrollo\b', texto_completo, re.IGNORECASE)
        if m_d:
            idx_d = m_d.start()
            inicio = texto_completo[:idx_d].strip()
            if "Inicio" in inicio or "Motivación" in inicio or "Saberes previos" in inicio:
                secuencia_didactica["inicio"] = inicio[-2000:] if len(inicio) > 2000 else inicio
            resto = texto_completo[idx_d:]
            m_c = re.search(r'\bCierre\b', resto, re.IGNORECASE)
            if m_c:
                idx_c = m_c.start()
                secuencia_didactica["desarrollo"] = (resto[:idx_c].strip())[:2000]
                secuencia_didactica["cierre"] = (resto[idx_c:].strip())[:2000]
            else:
                secuencia_didactica["desarrollo"] = resto.strip()[:2000]
    
    # Recursos por defecto si no se extraen (como en la imagen)
    recursos_inicio = "Diapositivas, Metaplan, Texto, Hoja de trabajo, Multimedia, Celular, Plataforma virtual, Kahoot, Tablets, Visualizador de material concreto, Video, Pizarras"
    recursos_desarrollo = "Interactivas, Afiches, Revistas, Periódicos, Post it, Plumones, Papelotes, Hojas de colores, Escuadras"
    if secuencia_didactica.get("recursos"):
        recursos_inicio = secuencia_didactica["recursos"][:500]
        recursos_desarrollo = secuencia_didactica["recursos"][:500]
    
    fecha_elaboracion = datetime.now().strftime('%d de %B de %Y')
    # Escapar HTML para evitar romper la tabla
    def esc(s):
        if not s:
            return ""
        return str(s).replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;").replace('"', "&quot;")
    
    # Tabla I. DATOS INFORMATIVOS (estructura de la imagen: 2 columnas iniciales, luego 4 columnas)
    tabla_datos_informativos = f"""
### I. DATOS INFORMATIVOS

<table style="width: 100%; border-collapse: collapse;">
<tr>
<td style="background-color: #2C3E50; color: #FFFFFF; padding: 8px; border: 1px solid #555; font-weight: bold; width: 22%;">Institución educativa</td>
<td colspan="3" style="background-color: #1E1E1E; color: #E0E0E0; padding: 8px; border: 1px solid #555;"></td>
</tr>
<tr>
<td style="background-color: #2C3E50; color: #FFFFFF; padding: 8px; border: 1px solid #555; font-weight: bold;">Título de la unidad</td>
<td colspan="3" style="background-color: #1E1E1E; color: #E0E0E0; padding: 8px; border: 1px solid #555;">{esc(titulo_unidad)}</td>
</tr>
<tr>
<td style="background-color: #2C3E50; color: #FFFFFF; padding: 8px; border: 1px solid #555; font-weight: bold; vertical-align: top;">Título de la sesión</td>
<td colspan="3" style="background-color: #1E1E1E; color: #E0E0E0; padding: 8px; border: 1px solid #555; vertical-align: top;">{esc(titulo_sesion)}</td>
</tr>
<tr>
<td style="background-color: #2C3E50; color: #FFFFFF; padding: 8px; border: 1px solid #555; font-weight: bold;">Nivel</td>
<td style="background-color: #1E1E1E; color: #E0E0E0; padding: 8px; border: 1px solid #555;">{esc(nivel)}</td>
<td style="background-color: #2C3E50; color: #FFFFFF; padding: 8px; border: 1px solid #555; font-weight: bold;">Área(s)</td>
<td style="background-color: #1E1E1E; color: #E0E0E0; padding: 8px; border: 1px solid #555;">{esc(area_curricular)}</td>
</tr>
<tr>
<td style="background-color: #2C3E50; color: #FFFFFF; padding: 8px; border: 1px solid #555; font-weight: bold;">Grado</td>
<td style="background-color: #1E1E1E; color: #E0E0E0; padding: 8px; border: 1px solid #555;">{esc(grado)}</td>
<td style="background-color: #2C3E50; color: #FFFFFF; padding: 8px; border: 1px solid #555; font-weight: bold;">Secciones</td>
<td style="background-color: #1E1E1E; color: #E0E0E0; padding: 8px; border: 1px solid #555;">{esc(seccion)}</td>
</tr>
<tr>
<td style="background-color: #2C3E50; color: #FFFFFF; padding: 8px; border: 1px solid #555; font-weight: bold;">Docente</td>
<td style="background-color: #1E1E1E; color: #E0E0E0; padding: 8px; border: 1px solid #555;"></td>
<td style="background-color: #2C3E50; color: #FFFFFF; padding: 8px; border: 1px solid #555; font-weight: bold;">Fecha</td>
<td style="background-color: #1E1E1E; color: #E0E0E0; padding: 8px; border: 1px solid #555;">{esc(fecha_elaboracion)}</td>
</tr>
<tr>
<td style="background-color: #2C3E50; color: #FFFFFF; padding: 8px; border: 1px solid #555; font-weight: bold;">Duración de la sesión</td>
<td colspan="3" style="background-color: #1E1E1E; color: #E0E0E0; padding: 8px; border: 1px solid #555;">{esc(duracion)}</td>
</tr>
</table>

---

### Situación significativa

<table style="width: 100%; border-collapse: collapse;">
<tr>
<td style="background-color: #2C3E50; color: #FFFFFF; padding: 8px; border: 1px solid #555; font-weight: bold; text-align: center;">Situación significativa</td>
</tr>
<tr>
<td style="background-color: #1E1E1E; color: #E0E0E0; padding: 8px; border: 1px solid #555;">{esc(situacion_significativa)}</td>
</tr>
</table>

---

### II. PROPÓSITOS DE APRENDIZAJE

<table style="width: 100%; border-collapse: collapse;">
<tr>
<td style="background-color: #2C3E50; color: #FFFFFF; padding: 8px; border: 1px solid #555; font-weight: bold; text-align: center;">Competencias</td>
<td style="background-color: #2C3E50; color: #FFFFFF; padding: 8px; border: 1px solid #555; font-weight: bold; text-align: center;">Capacidades</td>
<td style="background-color: #2C3E50; color: #FFFFFF; padding: 8px; border: 1px solid #555; font-weight: bold; text-align: center;">Criterios de evaluación</td>
<td style="background-color: #2C3E50; color: #FFFFFF; padding: 8px; border: 1px solid #555; font-weight: bold; text-align: center;">Contenidos</td>
<td style="background-color: #2C3E50; color: #FFFFFF; padding: 8px; border: 1px solid #555; font-weight: bold; text-align: center;">Evidencia de aprendizaje</td>
<td rowspan="7" style="background-color: #1E1E1E; color: #E0E0E0; padding: 8px; border: 1px solid #555; vertical-align: top; width: 14%;"><strong style="display: block; margin-bottom: 6px; color: #E0E0E0;">Instrumento de evaluación</strong>{esc(propositos_sesion.get('instrumento', ''))}</td>
</tr>
<tr>
<td style="background-color: #1E1E1E; color: #E0E0E0; padding: 8px; border: 1px solid #555;">{esc(propositos_sesion.get('competencias', ''))}</td>
<td style="background-color: #1E1E1E; color: #E0E0E0; padding: 8px; border: 1px solid #555;">{esc(propositos_sesion.get('capacidades', ''))}</td>
<td style="background-color: #1E1E1E; color: #E0E0E0; padding: 8px; border: 1px solid #555;">{esc(propositos_sesion.get('criterios', ''))}</td>
<td style="background-color: #1E1E1E; color: #E0E0E0; padding: 8px; border: 1px solid #555;">{esc(propositos_sesion.get('contenidos', ''))}</td>
<td style="background-color: #1E1E1E; color: #E0E0E0; padding: 8px; border: 1px solid #555;">{esc(propositos_sesion.get('evidencia', ''))}</td>
</tr>
<tr>
<td style="background-color: #2C3E50; color: #FFFFFF; padding: 8px; border: 1px solid #555; font-weight: bold;">Competencias transversales</td>
<td colspan="2" style="background-color: #2C3E50; color: #FFFFFF; padding: 8px; border: 1px solid #555; font-weight: bold;">Capacidad transversal</td>
<td colspan="2" style="background-color: #2C3E50; color: #FFFFFF; padding: 8px; border: 1px solid #555; font-weight: bold;">Desempeño transversal</td>
</tr>
<tr>
<td style="background-color: #1E1E1E; color: #E0E0E0; padding: 8px; border: 1px solid #555;">{esc(comp_trans_sesion[0]['competencia'])}</td>
<td colspan="2" style="background-color: #1E1E1E; color: #E0E0E0; padding: 8px; border: 1px solid #555;">{esc(comp_trans_sesion[0].get('capacidad', ''))}</td>
<td colspan="2" style="background-color: #1E1E1E; color: #E0E0E0; padding: 8px; border: 1px solid #555;">{esc(comp_trans_sesion[0].get('desempeno', ''))}</td>
</tr>
<tr>
<td style="background-color: #1E1E1E; color: #E0E0E0; padding: 8px; border: 1px solid #555;">{esc(comp_trans_sesion[1]['competencia'])}</td>
<td colspan="2" style="background-color: #1E1E1E; color: #E0E0E0; padding: 8px; border: 1px solid #555;">{esc(comp_trans_sesion[1].get('capacidad', ''))}</td>
<td colspan="2" style="background-color: #1E1E1E; color: #E0E0E0; padding: 8px; border: 1px solid #555;">{esc(comp_trans_sesion[1].get('desempeno', ''))}</td>
</tr>
<tr>
<td style="background-color: #2C3E50; color: #FFFFFF; padding: 8px; border: 1px solid #555; font-weight: bold;">Enfoque transversal</td>
<td colspan="2" style="background-color: #2C3E50; color: #FFFFFF; padding: 8px; border: 1px solid #555; font-weight: bold;">Valor priorizado / Valor operativo</td>
<td colspan="2" style="background-color: #2C3E50; color: #FFFFFF; padding: 8px; border: 1px solid #555; font-weight: bold;">Comportamientos observables</td>
</tr>
<tr>
<td style="background-color: #1E1E1E; color: #E0E0E0; padding: 8px; border: 1px solid #555;"></td>
<td colspan="2" style="background-color: #1E1E1E; color: #E0E0E0; padding: 8px; border: 1px solid #555;">{esc((enfoque_sesion.get('valor_priorizado') or '') + (' / ' if ((enfoque_sesion.get('valor_priorizado') or '') and (enfoque_sesion.get('valor_operativo') or '')) else '') + (enfoque_sesion.get('valor_operativo') or ''))}</td>
<td colspan="2" style="background-color: #1E1E1E; color: #E0E0E0; padding: 8px; border: 1px solid #555;">{esc(enfoque_sesion.get('comportamientos', ''))}</td>
</tr>
</table>

---

### III. SECUENCIA DIDÁCTICA

<table style="width: 100%; border-collapse: collapse;">
<tr>
<td style="background-color: #2C3E50; color: #FFFFFF; padding: 8px; border: 1px solid #555; font-weight: bold; text-align: center;">Procesos pedagógicos</td>
<td style="background-color: #2C3E50; color: #FFFFFF; padding: 8px; border: 1px solid #555; font-weight: bold; text-align: center;">Estrategias / Actividades / Procesos didácticos</td>
<td style="background-color: #2C3E50; color: #FFFFFF; padding: 8px; border: 1px solid #555; font-weight: bold; text-align: center;">Recursos y materiales</td>
</tr>
<tr>
<td style="background-color: #1E1E1E; color: #E0E0E0; padding: 8px; border: 1px solid #555; font-weight: bold; vertical-align: top; width: 12%;">Inicio</td>
<td style="background-color: #1E1E1E; color: #E0E0E0; padding: 8px; border: 1px solid #555; vertical-align: top;">{esc(secuencia_didactica.get('inicio', ''))}</td>
<td style="background-color: #1E1E1E; color: #E0E0E0; padding: 8px; border: 1px solid #555; vertical-align: top;">{esc(recursos_inicio)}</td>
</tr>
<tr>
<td style="background-color: #1E1E1E; color: #E0E0E0; padding: 8px; border: 1px solid #555; font-weight: bold; vertical-align: top;">Desarrollo</td>
<td style="background-color: #1E1E1E; color: #E0E0E0; padding: 8px; border: 1px solid #555; vertical-align: top;">{esc(secuencia_didactica.get('desarrollo', ''))}</td>
<td style="background-color: #1E1E1E; color: #E0E0E0; padding: 8px; border: 1px solid #555; vertical-align: top;">{esc(recursos_desarrollo)}</td>
</tr>
<tr>
<td style="background-color: #1E1E1E; color: #E0E0E0; padding: 8px; border: 1px solid #555; font-weight: bold; vertical-align: top;">Cierre</td>
<td style="background-color: #1E1E1E; color: #E0E0E0; padding: 8px; border: 1px solid #555; vertical-align: top;">{esc(secuencia_didactica.get('cierre', ''))}</td>
<td style="background-color: #1E1E1E; color: #E0E0E0; padding: 8px; border: 1px solid #555; vertical-align: top;"></td>
</tr>
</table>

---

### IV. REFLEXIONAMOS SOBRE LA ACTIVIDAD

<table style="width: 100%; border-collapse: collapse;">
<tr>
<td style="background-color: #2C3E50; color: #FFFFFF; padding: 8px; border: 1px solid #555; font-weight: bold; text-align: center;">¿Qué dificultades se observaron en el desarrollo de la sesión?</td>
</tr>
<tr>
<td style="background-color: #1E1E1E; color: #E0E0E0; padding: 8px; border: 1px solid #555; vertical-align: top; white-space: pre-wrap;">{esc(reflexion_actividad)}</td>
</tr>
</table>

---
"""
    
    contenido_formateado = f"""
# 📖 SESIÓN DE APRENDIZAJE

## 📚 {titulo_unidad}

### 🎯 {titulo_sesion}

---

{tabla_datos_informativos}

### 📝 NOTAS METODOLÓGICAS

Esta sesión de aprendizaje ha sido diseñada siguiendo los lineamientos del Currículo Nacional de la Educación Básica del Perú.

**Recomendaciones de implementación:**
- Considerar el contexto sociocultural de los estudiantes
- Adaptar las estrategias según los ritmos de aprendizaje
- Integrar recursos tecnológicos disponibles
- Promover el aprendizaje colaborativo

---

*Documento generado automáticamente por el Sistema de IA Educativa*
"""
    return contenido_formateado

# Función para obtener la ruta del Desktop
def obtener_ruta_desktop():
    """Obtiene la ruta del directorio Desktop del usuario"""
    # Si estamos en Docker, usar /app/outputs
    if os.path.exists("/app/outputs"):
        outputs_dir = Path("/app/outputs")
        outputs_dir.mkdir(parents=True, exist_ok=True)
        return outputs_dir
    
    # Si existe /app/desktop_outputs (montado desde Docker), usarlo
    if os.path.exists("/app/desktop_outputs"):
        outputs_dir = Path("/app/desktop_outputs")
        outputs_dir.mkdir(parents=True, exist_ok=True)
        return outputs_dir
    
    # Caso normal: usar Desktop del usuario
    home = Path.home()
    desktop = home / "Desktop"
    # Crear carpeta de outputs si no existe
    outputs_dir = desktop / "content_edu_outputs"
    outputs_dir.mkdir(parents=True, exist_ok=True)
    return outputs_dir

# Función para guardar archivo en Desktop
def guardar_archivo_desktop(contenido, nombre_archivo, es_bytes=False):
    """
    Guarda un archivo en el Desktop del usuario
    Args:
        contenido: Contenido del archivo (str o bytes)
        nombre_archivo: Nombre del archivo
        es_bytes: True si el contenido es bytes (para DOCX), False si es texto
    Returns:
        Ruta completa del archivo guardado
    """
    try:
        desktop_dir = obtener_ruta_desktop()
        ruta_completa = desktop_dir / nombre_archivo
        
        if es_bytes:
            with open(ruta_completa, 'wb') as f:
                f.write(contenido)
        else:
            with open(ruta_completa, 'w', encoding='utf-8') as f:
                f.write(contenido)
        
        return str(ruta_completa)
    except Exception as e:
        print(f"Error guardando archivo en Desktop: {e}")
        return None

def procesar_contenido_celda_tabla(celda, celda_word):
    """
    Procesa el contenido de una celda de tabla y formatea correctamente las viñetas y listas.
    
    Args:
        celda: Contenido de la celda como string (puede tener múltiples líneas y viñetas)
        celda_word: Objeto de celda de Word donde se insertará el contenido
    """
    if not celda or not celda.strip():
        return
    
    # Limpiar el texto primero (remover markdown bold)
    celda = celda.replace('**', '').strip()
    
    # Dividir por líneas
    lineas = celda.split('\n')
    
    # Limpiar la celda primero (eliminar el párrafo por defecto)
    if len(celda_word.paragraphs) > 0:
        celda_word.paragraphs[0].clear()
    else:
        celda_word.add_paragraph()
    
    # Procesar cada línea
    for idx, linea in enumerate(lineas):
        linea_original = linea
        linea = linea.strip()
        
        if not linea:
            # Si la línea está vacía, agregar un párrafo vacío solo si hay más líneas después
            if idx < len(lineas) - 1:
                celda_word.add_paragraph()
            continue
        
        # Detectar si es una viñeta
        es_viñeta = False
        texto_viñeta = linea
        
        # Verificar diferentes tipos de viñetas
        # Patrón 1: Viñetas comunes al inicio (•, -, *, →, ▪, ▫, ○, ●) con o sin espacios
        if re.match(r'^[\s]*[•\-\*→▪▫○●][\s]*', linea):
            es_viñeta = True
            # Remover el carácter de viñeta y espacios iniciales
            texto_viñeta = re.sub(r'^[\s]*[•\-\*→▪▫○●][\s]*', '', linea).strip()
        # Patrón 2: Lista numerada (1. , 1) , 1- )
        elif re.match(r'^[\s]*\d+[\.\)\-][\s]+', linea):
            es_viñeta = True
            # Mantener el número pero limpiar espacios extra al inicio
            texto_viñeta = re.sub(r'^[\s]+', '', linea)
        # Patrón 3: Viñetas simples sin espacio (solo el carácter)
        elif len(linea) > 1 and linea[0] in ['•', '-', '*', '→', '▪', '▫', '○', '●']:
            es_viñeta = True
            texto_viñeta = linea[1:].strip()
        
        # Crear o usar párrafo en la celda
        if idx == 0 and len(celda_word.paragraphs) > 0:
            # Usar el primer párrafo (ya existe después de clear)
            para = celda_word.paragraphs[0]
        else:
            # Crear nuevo párrafo
            para = celda_word.add_paragraph()
        
        # Si es viñeta, aplicar estilo de lista
        if es_viñeta and texto_viñeta:
            para.style = 'List Bullet'
            para.add_run(texto_viñeta)
        elif texto_viñeta:
            # Texto normal
            para.add_run(texto_viñeta)

# Función mejorada para crear Word
def crear_documento_profesional(contenido, titulo, subtitulo_extra=""):
    if not DOCX_OK:
        return None
    
    doc = Document()
    
    # Configurar propiedades del documento
    doc.core_properties.title = titulo
    doc.core_properties.author = "Sistema IA Educativa"
    doc.core_properties.subject = titulo
    
    # Título principal
    titulo_principal = doc.add_heading(titulo.upper(), 0)
    titulo_principal.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Subtítulo si existe
    if subtitulo_extra:
        subtitulo = doc.add_heading(subtitulo_extra, 1)
        subtitulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Si es Unidad Didáctica, agregar cuadro de DATOS INFORMATIVOS al inicio
    if "Unidad Didáctica" in titulo or "unidad didáctica" in titulo.lower():
        # Título de la sección
        doc.add_paragraph()  # Espacio antes
        datos_heading = doc.add_heading("I. DATOS INFORMATIVOS", 2)
        datos_heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
        
        # Crear tabla de DATOS INFORMATIVOS con 4 columnas
        tabla_datos = doc.add_table(rows=5, cols=4)
        tabla_datos.style = 'Light Grid Accent 1'
        
        # Configurar ancho de columnas
        for idx, col in enumerate(tabla_datos.columns):
            if idx % 2 == 0:  # Columnas de etiquetas (0 y 2)
                col.width = Inches(1.5)
            else:  # Columnas de contenido (1 y 3)
                col.width = Inches(2.0)
        
        # Función auxiliar para aplicar color de fondo a una celda
        def aplicar_fondo_celda(celda, color_hex):
            """Aplica color de fondo a una celda usando XML"""
            from docx.oxml import parse_xml
            from docx.oxml.ns import nsdecls
            
            shading = parse_xml(
                f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>'
            )
            tcPr = celda._element.tcPr
            if tcPr is None:
                tcPr = parse_xml(f'<w:tcPr {nsdecls("w")}/>')
                celda._element.insert(0, tcPr)
            tcPr.append(shading)
        
        # Color azul claro en formato hexadecimal
        color_fondo_hex = "D9E1F2"  # RGB(217, 225, 242) en hex
        
        # Fila 1: DRE / UGEL | espacio | Institución educativa | espacio
        fila1 = tabla_datos.rows[0]
        celda_dre = fila1.cells[0]
        celda_dre.text = "DRE / UGEL"
        celda_dre.paragraphs[0].runs[0].font.bold = True
        aplicar_fondo_celda(celda_dre, color_fondo_hex)
        # Celda 1 vacía (espacio)
        fila1.cells[1].text = ""
        # Celda 2: Institución educativa
        celda_inst = fila1.cells[2]
        celda_inst.text = "Institución educativa"
        celda_inst.paragraphs[0].runs[0].font.bold = True
        aplicar_fondo_celda(celda_inst, color_fondo_hex)
        # Celda 3 vacía (espacio)
        fila1.cells[3].text = ""
        
        # Fila 2: Área curricular | espacio | Grado y secciones | espacio
        fila2 = tabla_datos.rows[1]
        celda_area = fila2.cells[0]
        celda_area.text = "Área curricular"
        celda_area.paragraphs[0].runs[0].font.bold = True
        aplicar_fondo_celda(celda_area, color_fondo_hex)
        # Celda 1 vacía (espacio)
        fila2.cells[1].text = ""
        # Celda 2: Grado y secciones
        celda_grado = fila2.cells[2]
        celda_grado.text = "Grado y secciones"
        celda_grado.paragraphs[0].runs[0].font.bold = True
        aplicar_fondo_celda(celda_grado, color_fondo_hex)
        # Celda 3 vacía (espacio)
        fila2.cells[3].text = ""
        
        # Fila 3: Fecha de inicio | espacio | Fecha de término | espacio
        fila3 = tabla_datos.rows[2]
        celda_fecha_ini = fila3.cells[0]
        celda_fecha_ini.text = "Fecha de inicio"
        celda_fecha_ini.paragraphs[0].runs[0].font.bold = True
        aplicar_fondo_celda(celda_fecha_ini, color_fondo_hex)
        # Celda 1 vacía (espacio)
        fila3.cells[1].text = ""
        # Celda 2: Fecha de término
        celda_fecha_fin = fila3.cells[2]
        celda_fecha_fin.text = "Fecha de término"
        celda_fecha_fin.paragraphs[0].runs[0].font.bold = True
        aplicar_fondo_celda(celda_fecha_fin, color_fondo_hex)
        # Celda 3 vacía (espacio)
        fila3.cells[3].text = ""
        
        # Fila 4: Coordinación pedagógica (combinar celdas 0-1) | espacio (combinar celdas 2-3)
        fila4 = tabla_datos.rows[3]
        celda_coord = fila4.cells[0]
        celda_coord.text = "Coordinación pedagógica"
        celda_coord.paragraphs[0].runs[0].font.bold = True
        aplicar_fondo_celda(celda_coord, color_fondo_hex)
        # Combinar celdas 0 y 1 (label)
        celda_coord.merge(fila4.cells[1])
        # Combinar celdas 2 y 3 (contenido)
        fila4.cells[2].merge(fila4.cells[3])
        fila4.cells[2].text = ""
        
        # Fila 5: Docente(s) (combinar celdas 0-1) | espacio (combinar celdas 2-3)
        fila5 = tabla_datos.rows[4]
        celda_docente = fila5.cells[0]
        celda_docente.text = "Docente(s)"
        celda_docente.paragraphs[0].runs[0].font.bold = True
        aplicar_fondo_celda(celda_docente, color_fondo_hex)
        # Combinar celdas 0 y 1 (label)
        celda_docente.merge(fila5.cells[1])
        # Combinar celdas 2 y 3 (contenido)
        fila5.cells[2].merge(fila5.cells[3])
        fila5.cells[2].text = ""
        
        doc.add_paragraph()  # Espacio después de la tabla
    
    # Información del documento
    info_para = doc.add_paragraph()
    info_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    info_run = info_para.add_run(f"Fecha: {datetime.now().strftime('%d/%m/%Y')}\nGenerado por: IA Educativa")
    info_run.italic = True
    
    # Línea separadora
    doc.add_paragraph("=" * 80)
    
    # No mostrar "CONTENIDO DE LA UNIDAD DIDÁCTICA" en Word (oculto por decisión de producto)
    es_unidad_didactica = "Unidad Didáctica" in titulo or "unidad didáctica" in titulo.lower()
    
    # Procesar contenido línea por línea con formato mejorado (omitido para Unidad Didáctica)
    lineas = contenido.split('\n')
    i = 0
    while i < len(lineas):
        if es_unidad_didactica:
            break
        line = lineas[i].strip()
        if not line:
            i += 1
            continue
        
        # Detectar tablas (líneas que empiezan y terminan con | - formato markdown completo)
        # PRIORIDAD: Si tiene | al inicio y final, es una tabla
        if re.match(r'^\s*\|.*\|\s*$', line) and line.count('|') >= 2:
            # Intentar crear una tabla real
            filas_tabla = []
            j = i
            dentro_tabla = True
            ultima_fila_completa = None
            
            # Recopilar líneas consecutivas que parecen ser parte de una tabla
            while j < len(lineas) and dentro_tabla:
                current_line = lineas[j]
                current_line_stripped = current_line.strip()
                
                # Si la línea tiene | y empieza y termina con |, es parte de la tabla (formato markdown completo)
                if re.match(r'^\s*\|.*\|\s*$', current_line_stripped) and current_line_stripped.count('|') >= 2:
                    # Verificar si es una línea separadora de markdown (solo contiene |, -, :, espacios)
                    es_separador = re.match(r'^\s*\|[\s\-\:]+\|\s*$', current_line_stripped)
                    if not es_separador:
                        # Dividir por | y filtrar celdas vacías al inicio y final
                        partes = current_line.split('|')
                        # Limpiar cada celda
                        fila = [celda.strip() for celda in partes]
                        # Eliminar celdas vacías al inicio y final (formato markdown)
                        while fila and not fila[0]:
                            fila.pop(0)
                        while fila and not fila[-1]:
                            fila.pop()
                        
                        # Validar y corregir el formato: debe tener exactamente 2 columnas [ITEM, CONTENIDO]
                        if len(fila) == 0:
                            # Fila vacía, saltar
                            pass
                        elif len(fila) == 1:
                            # Solo una columna: verificar si es un ITEM o CONTENIDO
                            contenido_unico = fila[0]
                            
                            # CRITERIOS MÁS ESTRICTOS PARA DETECTAR ITEM
                            es_item = False
                            contenido_len = len(contenido_unico)
                            
                            # Si es muy largo (> 100 caracteres), definitivamente es CONTENIDO
                            if contenido_len > 100:
                                es_item = False
                            # Si es corto y está en mayúsculas, probablemente es ITEM
                            elif contenido_len < 50 and contenido_unico.isupper():
                                es_item = True
                            # Si empieza con ** y es corto, es ITEM
                            elif contenido_unico.startswith('**') and contenido_len < 80:
                                es_item = True
                            # Si contiene palabras clave de ITEMs y es corto
                            elif contenido_len < 80:
                                palabras_item = ['TÍTULO', 'SITUACIÓN', 'COMPETENCIA', 'CAPACIDAD', 
                                                'EVIDENCIA', 'INSTRUMENTO', 'VALOR', 'SECUENCIA', 
                                                'ENFOQUE', 'SESIÓN', 'MATERIAL', 'REFLEXIÓN', 'ESTÁNDAR',
                                                'DESEMPEÑO', 'PROPÓSITO', 'ORGANIZACIÓN', 'EVALUACIÓN']
                                contenido_upper = contenido_unico.upper()
                                tiene_palabra_clave = any(
                                    contenido_upper.startswith(palabra) or 
                                    f' {palabra}' in contenido_upper or
                                    f'{palabra} ' in contenido_upper
                                    for palabra in palabras_item
                                )
                                if tiene_palabra_clave:
                                    es_item = True
                            
                            # Si la última fila tenía ITEM sin CONTENIDO, este contenido debe ir a CONTENIDO
                            if (len(filas_tabla) > 0 and 
                                len(filas_tabla[-1]) >= 1 and 
                                filas_tabla[-1][0] and 
                                not filas_tabla[-1][1]):
                                # Agregar este contenido a CONTENIDO de la última fila
                                filas_tabla[-1][1] = contenido_unico
                                ultima_fila_completa = len(filas_tabla) - 1
                            elif es_item:
                                # Es un ITEM, agregar como [ITEM, ""]
                                filas_tabla.append([contenido_unico, ""])
                                ultima_fila_completa = len(filas_tabla) - 1
                            else:
                                # Es CONTENIDO, agregar como ["", CONTENIDO] o a la última fila si tenía ITEM
                                if (len(filas_tabla) > 0 and 
                                    len(filas_tabla[-1]) >= 1 and 
                                    filas_tabla[-1][0] and 
                                    not filas_tabla[-1][1]):
                                    filas_tabla[-1][1] = contenido_unico
                                    ultima_fila_completa = len(filas_tabla) - 1
                                else:
                                    filas_tabla.append(["", contenido_unico])
                                    ultima_fila_completa = len(filas_tabla) - 1
                        elif len(fila) >= 2:
                            # Tiene 2 o más columnas: tomar solo las primeras 2 [ITEM, CONTENIDO]
                            item = fila[0].strip()
                            contenido = ' '.join(fila[1:]).strip()  # Unir todas las columnas adicionales en contenido
                            filas_tabla.append([item, contenido])
                            ultima_fila_completa = len(filas_tabla) - 1
                else:
                    # Línea sin | - puede ser contenido multilínea dentro de la última celda
                    # Solo agregar si:
                    # 1. Ya tenemos al menos una fila de tabla
                    # 2. La línea no está vacía
                    # 3. La línea no es claramente el inicio de otra sección (encabezado, lista, etc.)
                    if (len(filas_tabla) > 0 and 
                        current_line_stripped and 
                        not current_line_stripped.startswith('#') and
                        not current_line_stripped.startswith(('•', '-', '*', '→')) and
                        (not current_line_stripped.isupper() or len(current_line_stripped) < 5)):
                        # Agregar este contenido a la última celda de la última fila (columna CONTENIDO = índice 1)
                        if ultima_fila_completa is not None and len(filas_tabla[ultima_fila_completa]) >= 1:
                            # Asegurar que la fila tenga al menos 2 columnas
                            while len(filas_tabla[ultima_fila_completa]) < 2:
                                filas_tabla[ultima_fila_completa].append("")
                            # Agregar a la columna CONTENIDO (índice 1, segunda columna)
                            contenido_actual = filas_tabla[ultima_fila_completa][1] if len(filas_tabla[ultima_fila_completa]) > 1 else ""
                            if contenido_actual:
                                filas_tabla[ultima_fila_completa][1] = contenido_actual + '\n' + current_line_stripped
                            else:
                                filas_tabla[ultima_fila_completa][1] = current_line_stripped
                    else:
                        # Esta línea claramente no es parte de la tabla
                        dentro_tabla = False
                        break
                j += 1
            
            # Si tenemos al menos 1 fila (puede ser solo encabezado), crear tabla
            if len(filas_tabla) >= 1:
                # Validar y normalizar: todas las filas deben tener exactamente 2 columnas (ITEM | CONTENIDO)
                num_cols = 2  # Forzar 2 columnas
                
                # Asegurar que todas las filas tengan exactamente 2 columnas [ITEM, CONTENIDO]
                filas_normalizadas = []
                for fila in filas_tabla:
                    # Normalizar a exactamente 2 columnas
                    if len(fila) == 0:
                        # Fila vacía, crear fila con dos celdas vacías
                        fila_normalizada = ["", ""]
                    elif len(fila) == 1:
                        # Solo una columna: determinar si es ITEM o CONTENIDO
                        contenido_unico = fila[0].strip()
                        # Detectar si es un ITEM (títulos comunes en mayúsculas o con **)
                        es_item = (
                            contenido_unico.isupper() or
                            contenido_unico.startswith('**') or
                            contenido_unico.startswith('*') or
                            (len(contenido_unico) < 60 and any(
                                palabra in contenido_unico.upper() 
                                for palabra in ['TÍTULO', 'SITUACIÓN', 'COMPETENCIA', 'CAPACIDAD', 
                                               'EVIDENCIA', 'INSTRUMENTO', 'VALOR', 'SECUENCIA', 
                                               'ENFOQUE', 'SESIÓN', 'MATERIAL', 'REFLEXIÓN']
                            ))
                        )
                        if es_item:
                            # Es un ITEM, colocar en columna izquierda
                            fila_normalizada = [contenido_unico, ""]
                        else:
                            # Es CONTENIDO, colocar en columna derecha (solo si la última fila tenía ITEM)
                            # Si la última fila normalizada tenía ITEM pero no CONTENIDO, agregar aquí
                            if filas_normalizadas and filas_normalizadas[-1][0] and not filas_normalizadas[-1][1]:
                                filas_normalizadas[-1][1] = contenido_unico
                                continue  # Ya se agregó a la fila anterior
                            else:
                                # Nueva fila con ITEM vacío y CONTENIDO
                                fila_normalizada = ["", contenido_unico]
                    elif len(fila) >= 2:
                        # Tiene 2 o más columnas: [ITEM, CONTENIDO]
                        item = fila[0].strip()
                        # Unir todas las columnas adicionales en CONTENIDO
                        contenido = ' '.join([c.strip() for c in fila[1:] if c.strip()]).strip()
                        fila_normalizada = [item, contenido]
                    else:
                        # Caso por defecto: dos celdas vacías
                        fila_normalizada = ["", ""]
                    
                    # Asegurar que siempre tenga exactamente 2 columnas
                    while len(fila_normalizada) < 2:
                        fila_normalizada.append("")
                    fila_normalizada = fila_normalizada[:2]  # Tomar solo las primeras 2
                    filas_normalizadas.append(fila_normalizada)
                
                # Crear tabla en Word
                tabla = doc.add_table(rows=len(filas_normalizadas), cols=num_cols)
                tabla.style = 'Light Grid Accent 1'
                
                # Configurar ancho de columnas (primera columna más estrecha para ITEM, segunda más ancha para CONTENIDO)
                if len(tabla.columns) >= 2:
                    tabla.columns[0].width = Inches(1.5)  # Columna ITEM: 1.5 pulgadas
                    tabla.columns[1].width = Inches(5.5)  # Columna CONTENIDO: 5.5 pulgadas
                
                # Llenar la tabla
                for row_idx, fila in enumerate(filas_normalizadas):
                    # Asegurar que siempre tengamos exactamente 2 columnas: [ITEM, CONTENIDO]
                    item_texto = fila[0] if len(fila) > 0 else ""
                    contenido_texto = fila[1] if len(fila) > 1 else ""
                    
                    # Columna 0: ITEM (izquierda)
                    celda_item = tabla.rows[row_idx].cells[0]
                    celda_item.vertical_alignment = WD_ALIGN_VERTICAL.TOP
                    procesar_contenido_celda_tabla(item_texto, celda_item)
                    for paragraph in celda_item.paragraphs:
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    
                    # Columna 1: CONTENIDO (derecha)
                    celda_contenido = tabla.rows[row_idx].cells[1]
                    celda_contenido.vertical_alignment = WD_ALIGN_VERTICAL.TOP
                    procesar_contenido_celda_tabla(contenido_texto, celda_contenido)
                    for paragraph in celda_contenido.paragraphs:
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    
                    # Hacer la primera fila en negrita (encabezados)
                    if row_idx == 0:
                        for paragraph in celda_item.paragraphs:
                            for run in paragraph.runs:
                                run.bold = True
                        for paragraph in celda_contenido.paragraphs:
                            for run in paragraph.runs:
                                run.bold = True
                
                i = j - 1  # Ajustar índice
            else:
                # Si no es una tabla válida, agregar como texto
                cleaned_line = line.replace('|', ' | ').strip()
                doc.add_paragraph(cleaned_line)
        
        # Detectar encabezados (solo si no es tabla)
        elif line.startswith('#'):
            level = line.count('#')
            text = line.replace('#', '').strip()
            if text:
                doc.add_heading(text, level=min(level, 3))
        
        # Detectar listas con bullets
        elif line.startswith(('•', '-', '*', '→')):
            para = doc.add_paragraph()
            para.style = 'List Bullet'
            para.add_run(line[1:].strip())
        
        # Detectar texto en mayúsculas (posibles títulos)
        elif line.isupper() and len(line) > 5 and not line.startswith(('COMPETENCIA', 'CAPACIDAD', 'CONTENIDO', 'DESEMPEÑO', 'CRITERIO', 'INSTRUMENTO')):
            doc.add_heading(line.title(), 2)
        
        # Detectar secciones importantes (COMPETENCIA, CAPACIDADES, etc.)
        elif any(palabra in line.upper() for palabra in ['COMPETENCIA', 'CAPACIDADES', 'CONTENIDOS', 'DESEMPEÑOS', 'CRITERIOS', 'INSTRUMENTOS', 'TRANSVERSALES', 'SESIONES']):
            if line.isupper() and len(line) > 5:
                doc.add_heading(line.title(), 2)
            else:
                doc.add_paragraph(line)
        
        # Texto normal
        else:
            if len(line) > 5:  # Agregar líneas con contenido significativo
                doc.add_paragraph(line)
        
        i += 1
    
    # Pie de página
    doc.add_page_break()
    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run = footer.add_run("GENERADO POR SISTEMA IA EDUCATIVA\nMinisterio de Educación - República del Perú")
    footer_run.italic = True
    
    # Convertir a bytes
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()

# Función específica para crear documento de sesión de aprendizaje
def crear_documento_sesion_aprendizaje(contenido, titulo_unidad, titulo_sesion, nivel, grado, seccion, duracion="", area_curricular=""):
    """
    Crea un documento Word para sesión de aprendizaje con título de unidad y título de sesión.
    Incluye la tabla I. DATOS INFORMATIVOS al inicio.
    
    Args:
        contenido: Contenido de la sesión de aprendizaje
        titulo_unidad: Título de la unidad didáctica
        titulo_sesion: Título de la sesión de aprendizaje
        nivel: Nivel educativo
        grado: Grado
        seccion: Sección
        duracion: Duración de la sesión (opcional)
        area_curricular: Área curricular (opcional)
    """
    if not DOCX_OK:
        return None
    
    doc = Document()
    
    # Configurar propiedades del documento
    doc.core_properties.title = f"Sesión de Aprendizaje: {titulo_sesion}"
    doc.core_properties.author = "Sistema IA Educativa"
    doc.core_properties.subject = f"Unidad: {titulo_unidad} - Sesión: {titulo_sesion}"
    
    # Título principal
    titulo_principal = doc.add_heading("SESIÓN DE APRENDIZAJE", 0)
    titulo_principal.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Título de la unidad
    if titulo_unidad:
        titulo_unidad_heading = doc.add_heading(f"Unidad: {titulo_unidad}", 1)
        titulo_unidad_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Título de la sesión
    if titulo_sesion:
        titulo_sesion_heading = doc.add_heading(f"Sesión: {titulo_sesion}", 2)
        titulo_sesion_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # I. DATOS INFORMATIVOS (tabla como en el formato oficial)
    doc.add_heading("I. DATOS INFORMATIVOS", level=2)
    tabla_datos = doc.add_table(rows=7, cols=4)
    tabla_datos.style = "Table Grid"
    fecha_str = datetime.now().strftime('%d/%m/%Y')
    def merge_cells_row(tabla, row_idx, start_col, end_col):
        if end_col > start_col:
            tabla.rows[row_idx].cells[start_col].merge(tabla.rows[row_idx].cells[end_col])
    # Fila 0: Institución educativa | (colspan 3)
    r0 = tabla_datos.rows[0].cells
    r0[0].text = "Institución educativa"
    merge_cells_row(tabla_datos, 0, 1, 3)
    r0[1].text = ""
    # Fila 1: Título de la unidad | (colspan 3)
    r1 = tabla_datos.rows[1].cells
    r1[0].text = "Título de la unidad"
    merge_cells_row(tabla_datos, 1, 1, 3)
    r1[1].text = titulo_unidad or ""
    # Fila 2: Título de la sesión | (colspan 3) solo el título escogido por el usuario
    r2 = tabla_datos.rows[2].cells
    r2[0].text = "Título de la sesión"
    merge_cells_row(tabla_datos, 2, 1, 3)
    r2[1].text = titulo_sesion or ""
    # Filas 3-5: 4 columnas
    r3 = tabla_datos.rows[3].cells
    r3[0].text, r3[1].text, r3[2].text, r3[3].text = "Nivel", nivel or "", "Área(s)", area_curricular or ""
    r4 = tabla_datos.rows[4].cells
    r4[0].text, r4[1].text, r4[2].text, r4[3].text = "Grado", str(grado) if grado else "", "Secciones", seccion or ""
    r5 = tabla_datos.rows[5].cells
    r5[0].text, r5[1].text, r5[2].text, r5[3].text = "Docente", "", "Fecha", fecha_str
    # Fila 6: Duración | (colspan 3)
    r6 = tabla_datos.rows[6].cells
    r6[0].text = "Duración de la sesión"
    merge_cells_row(tabla_datos, 6, 1, 3)
    r6[1].text = duracion or ""
    doc.add_paragraph()
    
    # Tabla Situación significativa (una columna: encabezado + contenido generado por la IA)
    def _extraer_situacion(cont):
        for linea in cont.split('\n'):
            l = linea.strip()
            if '|' in l and 'SITUACIÓN SIGNIFICATIVA' in l.upper():
                partes = [p.strip() for p in l.split('|')]
                while partes and not partes[0]:
                    partes.pop(0)
                while partes and not partes[-1]:
                    partes.pop()
                if len(partes) >= 2:
                    return partes[1].replace('**', '').strip()
                if len(partes) >= 3:
                    return partes[2].replace('**', '').strip()
        return ""
    situacion_txt = _extraer_situacion(contenido)
    doc.add_heading("Situación significativa", level=2)
    tabla_sit = doc.add_table(rows=2, cols=1)
    tabla_sit.style = "Table Grid"
    tabla_sit.rows[0].cells[0].text = "Situación significativa"
    tabla_sit.rows[1].cells[0].text = situacion_txt or ""
    doc.add_paragraph()
    
    # Tabla II. PROPÓSITOS DE APRENDIZAJE (extraer datos del contenido generado)
    def _extraer_propositos(cont):
        d = {"competencias": "", "capacidades": "", "criterios": "", "contenidos": "", "evidencia": "", "instrumento": ""}
        for linea in cont.split('\n'):
            l = linea.strip()
            if '|' in l and 'PROPÓSITOS DE APRENDIZAJE' in l.upper():
                partes = [p.strip() for p in l.split('|')]
                while partes and not partes[0]: partes.pop(0)
                while partes and not partes[-1]: partes.pop()
                if len(partes) < 2: continue
                texto = partes[1].replace('**', '')
                for clave, sep in [("competencias", "Competencias:"), ("capacidades", "Capacidades:"), ("criterios", "Criterios de evaluación:"), ("contenidos", "Contenidos:"), ("evidencia", "Evidencia de aprendizaje:"), ("instrumento", "Instrumento de evaluación:")]:
                    if sep in texto:
                        idx = texto.find(sep)
                        resto = texto[idx + len(sep):].strip()
                        for s in ["Competencias:", "Capacidades:", "Criterios de evaluación:", "Contenidos:", "Evidencia de aprendizaje:", "Instrumento de evaluación:"]:
                            if s != sep and s in resto:
                                resto = resto[:resto.find(s)]
                                break
                        d[clave] = resto.strip()[:800]
                break
        return d
    def _extraer_comp_trans(cont):
        comp1, comp2 = "Se desenvuelve en los entornos virtuales generados por las TIC.", "Gestiona su aprendizaje de manera autónoma."
        res = [{"competencia": comp1, "capacidad": "", "desempeno": ""}, {"competencia": comp2, "capacidad": "", "desempeno": ""}]
        for linea in cont.split('\n'):
            l = linea.strip()
            if '|' in l and 'COMPETENCIAS TRANSVERSALES' in l.upper():
                partes = [p.strip() for p in l.split('|')]
                while partes and not partes[0]: partes.pop(0)
                while partes and not partes[-1]: partes.pop()
                if len(partes) >= 2:
                    texto = partes[1].replace('**', '')
                    bloques = re.split(re.escape(comp2), texto, 1)
                    for i, bloque in enumerate([bloques[0]] if len(bloques) == 1 else bloques):
                        if i == 1 and len(bloques) > 1: bloque = comp2 + bloque
                        if "Capacidad transversal:" in bloque:
                            resto = bloque.split("Capacidad transversal:", 1)[1]
                            if "Desempeño transversal:" in resto:
                                res[i]["capacidad"] = resto.split("Desempeño transversal:")[0].strip()[:400]
                                res[i]["desempeno"] = resto.split("Desempeño transversal:")[1].strip()[:500]
                            else: res[i]["capacidad"] = resto.strip()[:400]
                        if i == 0 and len(bloques) == 1:
                            res[1]["capacidad"], res[1]["desempeno"] = res[0]["capacidad"], res[0]["desempeno"]
                            break
                break
        return res
    def _extraer_enfoque(cont):
        d = {"valor_priorizado": "", "valor_operativo": "", "comportamientos": ""}
        for linea in cont.split('\n'):
            l = linea.strip()
            if '|' in l and 'ENFOQUE TRANSVERSAL' in l.upper():
                partes = [p.strip() for p in l.split('|')]
                while partes and not partes[0]: partes.pop(0)
                while partes and not partes[-1]: partes.pop()
                if len(partes) >= 2:
                    t = partes[1].replace('**', '')
                    if "Valor priorizado:" in t: d["valor_priorizado"] = t.split("Valor priorizado:")[1].split("Valor operativo:")[0].strip()[:300]
                    if "Valor operativo:" in t: d["valor_operativo"] = t.split("Valor operativo:")[1].split("Comportamientos observables:")[0].strip()[:300]
                    if "Comportamientos observables:" in t: d["comportamientos"] = t.split("Comportamientos observables:")[1].strip()[:600]
                break
        return d
    prop = _extraer_propositos(contenido)
    comp_trans = _extraer_comp_trans(contenido)
    enfoque = _extraer_enfoque(contenido)
    doc.add_heading("II. PROPÓSITOS DE APRENDIZAJE", level=2)
    t = doc.add_table(rows=6, cols=6)
    t.style = "Table Grid"
    headers1 = ["Competencias", "Capacidades", "Criterios de evaluación", "Contenidos", "Evidencia de aprendizaje", "Instrumento de evaluación"]
    for col, h in enumerate(headers1):
        t.rows[0].cells[col].text = h
    for col, k in enumerate(["competencias", "capacidades", "criterios", "contenidos", "evidencia", "instrumento"]):
        t.rows[1].cells[col].text = prop.get(k, "")
    t2 = doc.add_table(rows=4, cols=6)
    t2.style = "Table Grid"
    t2.rows[0].cells[0].text = "Competencias transversales"
    t2.rows[0].cells[1].merge(t2.rows[0].cells[2])
    t2.rows[0].cells[1].text = "Capacidad transversal"
    t2.rows[0].cells[3].merge(t2.rows[0].cells[4]).merge(t2.rows[0].cells[5])
    t2.rows[0].cells[3].text = "Desempeño transversal"
    t2.rows[1].cells[0].text = comp_trans[0]["competencia"]
    t2.rows[1].cells[1].merge(t2.rows[1].cells[2])
    t2.rows[1].cells[1].text = comp_trans[0].get("capacidad", "")
    t2.rows[1].cells[3].merge(t2.rows[1].cells[4]).merge(t2.rows[1].cells[5])
    t2.rows[1].cells[3].text = comp_trans[0].get("desempeno", "")
    t2.rows[2].cells[0].text = comp_trans[1]["competencia"]
    t2.rows[2].cells[1].merge(t2.rows[2].cells[2])
    t2.rows[2].cells[1].text = comp_trans[1].get("capacidad", "")
    t2.rows[2].cells[3].merge(t2.rows[2].cells[4]).merge(t2.rows[2].cells[5])
    t2.rows[2].cells[3].text = comp_trans[1].get("desempeno", "")
    t2.rows[3].cells[0].text = "Enfoque transversal"
    t2.rows[3].cells[1].merge(t2.rows[3].cells[2])
    t2.rows[3].cells[1].text = ((enfoque.get("valor_priorizado") or "") + (" / " if (enfoque.get("valor_priorizado") and enfoque.get("valor_operativo")) else "") + (enfoque.get("valor_operativo") or ""))
    t2.rows[3].cells[3].merge(t2.rows[3].cells[4]).merge(t2.rows[3].cells[5])
    t2.rows[3].cells[3].text = enfoque.get("comportamientos", "")
    doc.add_paragraph()
    
    # III. SECUENCIA DIDÁCTICA (en Word puede ir en una o dos tablas para que quepa en hoja)
    def _extraer_secuencia_didactica(cont):
        data = {"inicio": "", "desarrollo": "", "cierre": ""}
        contenido_sec = []
        dentro = False
        for linea in cont.split('\n'):
            l = linea.strip()
            if '|' in l and 'SECUENCIA DIDÁCTICA' in l.upper():
                dentro = True
                partes = [p.strip() for p in l.split('|')]
                while partes and not partes[0]: partes.pop(0)
                while partes and not partes[-1]: partes.pop()
                celda = (partes[2] if len(partes) >= 3 else partes[1] if len(partes) >= 2 else "").replace('**', '')
                if celda:
                    contenido_sec.append(celda)
                continue
            if dentro and '|' in l and '**' in l:
                item = (l.split('|')[1] or "").strip()
                if item and 'SECUENCIA' not in item.upper():
                    break
                partes = [p.strip() for p in l.split('|')]
                if len(partes) >= 3 and partes[2]:
                    contenido_sec.append(partes[2].replace('**', ''))
                elif len(partes) >= 2 and partes[1]:
                    contenido_sec.append(partes[1].replace('**', ''))
            elif dentro and l and not l.startswith('|'):
                if contenido_sec:
                    contenido_sec[-1] += "\n" + l
                else:
                    contenido_sec.append(l)
        texto = " ".join(contenido_sec).replace("\n", " ")
        if "Desarrollo" in texto:
            idx_d = texto.find("Desarrollo")
            data["inicio"] = texto[:idx_d].strip()
            resto = texto[idx_d:]
            if "Cierre" in resto:
                idx_c = resto.find("Cierre")
                data["desarrollo"] = resto[:idx_c].strip()
                data["cierre"] = resto[idx_c:].strip()
            else:
                data["desarrollo"] = resto.strip()
        else:
            data["inicio"] = texto.strip()
        return data
    sec_did = _extraer_secuencia_didactica(contenido)
    doc.add_heading("III. SECUENCIA DIDÁCTICA", level=2)
    # Tabla 1: Inicio (permite que en Word sea una tabla y si no cabe, la siguiente va a otra hoja)
    tab_sec1 = doc.add_table(rows=2, cols=3)
    tab_sec1.style = "Table Grid"
    tab_sec1.rows[0].cells[0].text = "Procesos pedagógicos"
    tab_sec1.rows[0].cells[1].text = "Estrategias / Actividades / Procesos didácticos"
    tab_sec1.rows[0].cells[2].text = "Recursos y materiales"
    tab_sec1.rows[1].cells[0].text = "Inicio"
    tab_sec1.rows[1].cells[1].text = sec_did.get("inicio", "")
    tab_sec1.rows[1].cells[2].text = "Diapositivas, Metaplan, Texto, Hoja de trabajo, Multimedia, Celular, Plataforma virtual, Kahoot, Tablets, Video, Pizarras"
    doc.add_paragraph()
    # Tabla 2: Desarrollo y Cierre (segunda tabla para que en Word pueda pasar a otra hoja si hace falta)
    tab_sec2 = doc.add_table(rows=3, cols=3)
    tab_sec2.style = "Table Grid"
    tab_sec2.rows[0].cells[0].text = "Procesos pedagógicos"
    tab_sec2.rows[0].cells[1].text = "Estrategias / Actividades / Procesos didácticos"
    tab_sec2.rows[0].cells[2].text = "Recursos y materiales"
    tab_sec2.rows[1].cells[0].text = "Desarrollo"
    tab_sec2.rows[1].cells[1].text = sec_did.get("desarrollo", "")
    tab_sec2.rows[1].cells[2].text = "Interactivas, Afiches, Revistas, Periódicos, Post it, Plumones, Papelotes, Hojas de colores, Escuadras"
    tab_sec2.rows[2].cells[0].text = "Cierre"
    tab_sec2.rows[2].cells[1].text = sec_did.get("cierre", "")
    tab_sec2.rows[2].cells[2].text = ""
    doc.add_paragraph()
    
    # IV. REFLEXIONAMOS SOBRE LA ACTIVIDAD (1 columna, 2 filas: encabezado + contenido con Dificultades, Mejoras, Ajustes)
    def _extraer_reflexion(cont):
        contenido_reflex = []
        dentro = False
        for linea in cont.split('\n'):
            l = linea.strip()
            if '|' in l and 'REFLEXIÓN' in l.upper() and 'ACTIVIDAD' in l.upper():
                dentro = True
                partes = [p.strip() for p in l.split('|')]
                while partes and not partes[0]: partes.pop(0)
                while partes and not partes[-1]: partes.pop()
                celda = (partes[2] if len(partes) >= 3 else partes[1] if len(partes) >= 2 else "").replace('**', '')
                if celda:
                    contenido_reflex.append(celda)
                continue
            if dentro and '|' in l:
                partes = [p.strip() for p in l.split('|')]
                primera_col = partes[1] if len(partes) > 1 else ""
                if primera_col and '**' in primera_col and 'REFLEXIÓN' not in primera_col.upper():
                    break
                contenido_celda = (partes[2] if len(partes) > 2 and partes[2] else partes[1] if len(partes) > 1 else "").replace('**', '')
                if contenido_celda:
                    contenido_reflex.append(contenido_celda)
            elif dentro and l and not l.startswith('|'):
                if contenido_reflex:
                    contenido_reflex[-1] += "\n" + l
                else:
                    contenido_reflex.append(l)
        return "\n\n".join(contenido_reflex).strip() if contenido_reflex else ""
    reflex_txt = _extraer_reflexion(contenido)
    doc.add_heading("IV. REFLEXIONAMOS SOBRE LA ACTIVIDAD", level=2)
    tab_reflex = doc.add_table(rows=2, cols=1)
    tab_reflex.style = "Table Grid"
    tab_reflex.rows[0].cells[0].text = "¿Qué dificultades se observaron en el desarrollo de la sesión?"
    tab_reflex.rows[1].cells[0].text = reflex_txt
    doc.add_paragraph()
    
    # Línea separadora
    doc.add_paragraph("=" * 80)
    
    # No mostrar "CONTENIDO DE LA SESIÓN DE APRENDIZAJE" en Word (oculto por decisión de producto)
    # El documento termina aquí: I. DATOS, Situación, II. PROPÓSITOS, III. SECUENCIA DIDÁCTICA, IV. REFLEXIONAMOS
    
    # Pie de página
    doc.add_page_break()
    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_text = f"GENERADO POR SISTEMA IA EDUCATIVA\n"
    footer_text += f"Unidad: {titulo_unidad}\n"
    footer_text += f"Sesión: {titulo_sesion}\n"
    footer_text += "Ministerio de Educación - República del Perú"
    footer_run = footer.add_run(footer_text)
    footer_run.italic = True
    
    # Convertir a bytes
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()

# Solo mostrar tabs si todo está OK
if SERVICES_OK:
    st.success("🎉 ¡Sistema listo! Genera tu contenido educativo.")
    
    # Crear tabs
    tab1, tab2 = st.tabs(["📚 Unidad Didáctica", "📖 Sesión de Aprendizaje"])
    
    with tab1:
        st.header("📚 Generador de Unidad Didáctica")
        
        # Inicializar chat desde el inicio
        if 'chat_mensajes_unidad' not in st.session_state:
            st.session_state['chat_mensajes_unidad'] = []
        
        # Información contextual
        with st.expander("ℹ️ Información sobre la Unidad Didáctica"):
            st.markdown("""
            **¿Qué incluye una unidad didáctica?**
            - ✅ Competencias y capacidades específicas
            - ✅ Contenidos organizados por temas
            - ✅ Desempeños observables y medibles
            - ✅ Criterios e instrumentos de evaluación
            - ✅ Estrategias metodológicas
            
            **Basado en:** Currículo Nacional de Educación Básica - MINEDU Perú
            """)
        
        with st.form("form_unidad", clear_on_submit=False):
            # Menús según malla curricular (Programa Curricular Educación Secundaria – Perú 2016)
            # Opción en blanco por defecto; el resto son sugerencias
            lista_areas = obtener_areas_curriculares_secundaria() if COMPETENCIAS_DISPONIBLES else AREAS_MALLA_FALLBACK
            lista_grados = obtener_grados_secundaria() if COMPETENCIAS_DISPONIBLES else GRADOS_MALLA_FALLBACK
            opciones_areas = ["— Seleccione un área curricular —"] + list(lista_areas)
            opciones_grados = ["— Seleccione un grado —"] + list(lista_grados)

            area_curricular = st.selectbox(
                "📚 Área Curricular",
                options=opciones_areas,
                index=0,
                help="Selecciona el área curricular según la malla de Educación Secundaria"
            )
            grado = st.selectbox(
                "🎓 Grado / Curso",
                options=opciones_grados,
                index=0,
                help="Selecciona el grado (curso) de secundaria: 1° a 5°"
            )

            # Selector de competencias con checkboxes (opcional, solo si está disponible)
            # Inicializar session_state para mantener las competencias seleccionadas
            if 'competencias_seleccionadas_unidad' not in st.session_state:
                st.session_state['competencias_seleccionadas_unidad'] = []
            if 'area_curricular_anterior' not in st.session_state:
                st.session_state['area_curricular_anterior'] = ""
            
            # Limpiar competencias si cambió el área curricular
            area_actual = area_curricular.strip() if area_curricular and area_curricular.strip() and not area_curricular.startswith("— Seleccione") else ""
            area_anterior = st.session_state.get('area_curricular_anterior', '')
            
            if area_actual and area_actual != area_anterior:
                # Limpiar competencias cuando cambia el área
                st.session_state['competencias_seleccionadas_unidad'] = []
                st.session_state['area_curricular_anterior'] = area_actual
            elif not area_actual:
                # Si no hay área seleccionada, limpiar también
                st.session_state['competencias_seleccionadas_unidad'] = []
            
            competencias_seleccionadas = []
            if COMPETENCIAS_DISPONIBLES:
                try:
                    competencias_relacionadas = []
                    if area_curricular and area_curricular.strip():
                        competencias_relacionadas = obtener_competencias_por_area(area_curricular.strip())
                    
                    if competencias_relacionadas:
                        competencias_opciones = [formatear_competencia_para_tabla(comp) for comp in competencias_relacionadas]
                        # Filtrar competencias seleccionadas previas que aún están disponibles
                        competencias_previas = st.session_state.get('competencias_seleccionadas_unidad', [])
                        competencias_validas = [c for c in competencias_previas if c in competencias_opciones]
                        # Usar session_state para mantener la selección válida
                        # Asegurar que default sea una lista válida
                        default_competencias = competencias_validas if competencias_validas else []
                        competencias_seleccionadas = st.multiselect(
                            "🎯 Competencias",
                            options=competencias_opciones,
                            help="Selecciona una o más competencias relacionadas con el área curricular",
                            default=default_competencias,
                            key="multiselect_competencias_unidad_area"
                        )
                        # Actualizar session_state con la selección actual siempre
                        st.session_state['competencias_seleccionadas_unidad'] = list(competencias_seleccionadas) if competencias_seleccionadas else []
                    else:
                        # Si no hay área o no se encontraron competencias, mostrar todas
                        todas_competencias = obtener_todas_las_competencias()
                        if todas_competencias:
                            competencias_opciones = [formatear_competencia_para_tabla(comp) for comp in todas_competencias]
                            # Filtrar competencias seleccionadas previas que aún están disponibles
                            competencias_previas = st.session_state.get('competencias_seleccionadas_unidad', [])
                            competencias_validas = [c for c in competencias_previas if c in competencias_opciones]
                            # Asegurar que default sea una lista válida
                            default_competencias = competencias_validas if competencias_validas else []
                            competencias_seleccionadas = st.multiselect(
                                "🎯 Competencias",
                                options=competencias_opciones,
                                help="Selecciona una o más competencias del Currículo Nacional",
                                default=default_competencias,
                                key="multiselect_competencias_unidad_todas"
                            )
                            # Actualizar session_state con la selección actual
                            st.session_state['competencias_seleccionadas_unidad'] = competencias_seleccionadas if competencias_seleccionadas else []
                except Exception:
                    # Si hay algún error, simplemente no mostrar el selector
                    competencias_seleccionadas = []
                    st.session_state['competencias_seleccionadas_unidad'] = []
            
            # Campo para temas
            temas_unidad = st.text_area(
                "📝 Temas (opcional)",
                help="Especifica los temas o contenidos que deseas incluir en la unidad didáctica",
                placeholder="Ejemplo: Temas relacionados con el área curricular seleccionada...",
                height=80
            )
            
            # Campo para número de sesiones
            num_sesiones = st.number_input(
                "🔢 Número de sesiones",
                min_value=4,
                value=6,
                step=1,
                help="Especifica cuántas sesiones de aprendizaje tendrá la unidad didáctica (mínimo 4)"
            )
            
            generar = st.form_submit_button("🎯 Generar Unidad Didáctica", use_container_width=True)
        
        # FUERA del formulario - manejar resultados
        if generar:
            # Tratar la opción por defecto (placeholder) como no seleccionado
            area_vacia = not area_curricular.strip() or area_curricular.startswith("— Seleccione")
            grado_vacio = not grado.strip() or grado.startswith("— Seleccione")
            if area_vacia:
                st.warning("⚠️ Por favor selecciona un área curricular")
            elif grado_vacio:
                st.warning("⚠️ Por favor selecciona un grado")
            else:
                with st.spinner('🔄 Generando unidad didáctica...'):
                    try:
                        # Pasar las competencias seleccionadas si existen y están disponibles
                        competencia_para_generar = None
                        competencias_para_usar = st.session_state.get('competencias_seleccionadas_unidad', [])
                        if COMPETENCIAS_DISPONIBLES and competencias_para_usar:
                            # Si hay múltiples competencias, concatenarlas con saltos de línea
                            if len(competencias_para_usar) > 0:
                                competencia_para_generar = "\n".join(competencias_para_usar)
                        
                        # Obtener temas y número de sesiones del formulario
                        temas_para_generar = temas_unidad.strip() if temas_unidad and temas_unidad.strip() else None
                        num_sesiones_para_generar = num_sesiones if num_sesiones >= 4 else 4
                        
                        resultado_raw = generar_unidad_didactica(
                            area_curricular, 
                            grado, 
                            competencia_referencia=competencia_para_generar,
                            temas=temas_para_generar,
                            num_sesiones=num_sesiones_para_generar
                        )
                        
                        # Formatear el contenido con encabezados y estructura completa (retorna también títulos de sesiones de la tabla IV)
                        contenido_formateado, titulos_sesiones = formatear_unidad_didactica(resultado_raw, area_curricular, num_sesiones_para_generar)
                        
                        # Guardar archivos automáticamente en Desktop
                        fecha_str = datetime.now().strftime('%Y%m%d_%H%M%S')
                        nombre_txt = f"unidad_didactica_{fecha_str}.txt"
                        ruta_txt = guardar_archivo_desktop(contenido_formateado, nombre_txt, es_bytes=False)
                        
                        if DOCX_OK:
                            doc_bytes = crear_documento_profesional(resultado_raw, "Unidad Didáctica", f"Área: {area_curricular}")
                            if doc_bytes:
                                nombre_docx = f"unidad_didactica_{fecha_str}.docx"
                                ruta_docx = guardar_archivo_desktop(doc_bytes, nombre_docx, es_bytes=True)
                        else:
                            doc_bytes = None
                            ruta_docx = None
                        
                        # Título de la unidad didáctica (títulos de sesiones ya vienen de formatear_unidad_didactica)
                        titulo_unidad = extraer_titulo_unidad_didactica(resultado_raw)
                        if not titulo_unidad:
                            titulo_unidad = f"Unidad Didáctica - {area_curricular}"
                        
                        # Guardar en sesión para usar en sesión de aprendizaje
                        st.session_state['unidad_generada'] = {
                            'titulo': titulo_unidad,
                            'area_curricular': area_curricular,
                            'grado': grado,
                            'contenido': resultado_raw,
                            'titulos_sesiones': titulos_sesiones
                        }
                        # Guardar documento editable y reiniciar chat de mejoras
                        st.session_state['documento_editable_unidad'] = contenido_formateado
                        st.session_state['documento_raw_unidad'] = resultado_raw
                        st.session_state['chat_mensajes_unidad'] = []
                        
                        st.success("✅ ¡Unidad didáctica generada exitosamente!")
                        if ruta_txt:
                            st.info(f"📁 Archivos guardados en: {ruta_txt.rsplit('/', 1)[0]}")
                                
                    except Exception as e:
                        st.error(f"❌ Error generando unidad didáctica: {str(e)}")
                        st.info("💡 Verifica la conexión con AWS Bedrock")
        
        # Mostrar documento actual (generado o mejorado por chat) y chat de mejoras
        if st.session_state.get('documento_editable_unidad'):
            st.markdown("---")
            
            # Sección de documento con expander para mejor organización
            with st.expander("📄 Ver documento actual", expanded=True):
                doc_actual = st.session_state['documento_editable_unidad']
                st.markdown(doc_actual, unsafe_allow_html=True)
            
            # Botones de acción en una fila organizada
            st.markdown("### 📥 Descargar documento")
            col1, col2, col3 = st.columns([2, 2, 2])
            with col1:
                st.download_button(
                    "📄 Descargar TXT",
                    data=doc_actual,
                    file_name=f"unidad_didactica_{datetime.now().strftime('%Y%m%d_%H%M%S')}.txt",
                    mime="text/plain",
                    key="download_txt_unidad",
                    use_container_width=True
                )
            with col2:
                if DOCX_OK:
                    doc_bytes_actual = crear_documento_profesional(
                        st.session_state.get('documento_raw_unidad', doc_actual),
                        "Unidad Didáctica",
                        f"Área: {st.session_state.get('unidad_generada', {}).get('area_curricular', '')}"
                    )
                    if doc_bytes_actual:
                        st.download_button(
                            "📝 Descargar WORD",
                            data=doc_bytes_actual,
                            file_name=f"unidad_didactica_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx",
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                            key="download_docx_unidad",
                            use_container_width=True
                        )
                    else:
                        st.button("📝 WORD no disponible", disabled=True, key="docx_disabled_unidad", use_container_width=True)
                else:
                    st.button("📝 WORD no disponible", disabled=True, key="docx_disabled_unidad", use_container_width=True)
            with col3:
                if st.button("🔄 Generar Nueva Unidad", key="nueva_unidad", use_container_width=True):
                    for k in ('documento_editable_unidad', 'documento_raw_unidad', 'chat_mensajes_unidad'):
                        if k in st.session_state:
                            del st.session_state[k]
                    st.rerun()
            
            st.markdown("---")
        
        # Chat siempre visible desde el inicio - Mejor organizado
        st.markdown("### 💬 Editor de Chat - Mejorar documento")
        
        # Mostrar estado del documento
        if st.session_state.get('documento_editable_unidad'):
            st.success("✅ Tienes un documento generado. Puedes mejorarlo usando el chat.")
        else:
            st.info("ℹ️ **Genera primero una unidad didáctica arriba para poder mejorarla con el chat.**")
        
        st.info("💡 **Sugerencias:** Puedes pedir cambios como 'haz más breve la sección de criterios', 'mejora el lenguaje', 'añade más ejemplos', etc.")
        
        # Contenedor para el chat con mejor estilo
        chat_container = st.container()
        with chat_container:
            # Mostrar historial de chat con mejor formato
            if st.session_state['chat_mensajes_unidad']:
                st.markdown("#### 📜 Historial de conversación")
                for idx, msg in enumerate(st.session_state['chat_mensajes_unidad']):
                    with st.chat_message(msg["role"]):
                        if msg["role"] == "user":
                            st.markdown(f"**Tu solicitud:**\n{msg['content']}")
                        else:
                            # Mejorar formato de respuesta del asistente
                            contenido = msg['content']
                            if "✅ Cambios aplicados" in contenido:
                                st.success("✅ **Cambios aplicados exitosamente**")
                                # Extraer solo la vista previa si existe
                                if "Vista previa:" in contenido:
                                    partes = contenido.split("Vista previa:", 1)
                                    if len(partes) > 1:
                                        st.markdown(f"**Vista previa:**\n{partes[1].strip()}")
                            elif "❌ Error" in contenido:
                                st.error(contenido)
                            else:
                                st.markdown(contenido)
                    if idx < len(st.session_state['chat_mensajes_unidad']) - 1:
                        st.markdown("---")
            else:
                st.markdown("*No hay mensajes aún. Escribe abajo para comenzar a mejorar el documento.*")
            
            # Input de chat con mejor placeholder (habilitado solo si hay documento)
            tiene_documento = st.session_state.get('documento_editable_unidad')
            prompt_chat = st.chat_input(
                "Escribe aquí cómo quieres mejorar el documento..." if tiene_documento else "Primero genera una unidad didáctica arriba...",
                key="chat_input_unidad",
                disabled=not tiene_documento
            )
            
            if prompt_chat and tiene_documento:
                # Agregar mensaje del usuario al historial
                st.session_state['chat_mensajes_unidad'].append({"role": "user", "content": prompt_chat})
                
                # Mostrar spinner mientras se procesa
                with st.spinner("🔄 Aplicando cambios al documento..."):
                    try:
                        nuevo_doc = mejorar_documento_con_instruccion(
                            st.session_state['documento_editable_unidad'],
                            prompt_chat,
                            "unidad didáctica"
                        )
                        
                        if nuevo_doc and not nuevo_doc.startswith("[Error"):
                            st.session_state['documento_editable_unidad'] = nuevo_doc
                            st.session_state['documento_raw_unidad'] = nuevo_doc
                            # Mensaje de éxito más claro
                            st.session_state['chat_mensajes_unidad'].append({
                                "role": "assistant",
                                "content": f"✅ Cambios aplicados exitosamente. El documento ha sido actualizado.\n\n**Vista previa de los cambios:**\n\n{nuevo_doc[:400]}..."
                            })
                        else:
                            st.session_state['chat_mensajes_unidad'].append({
                                "role": "assistant",
                                "content": f"⚠️ No se pudieron aplicar los cambios. Por favor, intenta con una instrucción más específica."
                            })
                    except Exception as e:
                        st.session_state['chat_mensajes_unidad'].append({
                            "role": "assistant",
                            "content": f"❌ Error al procesar la solicitud: {str(e)}\n\nPor favor, intenta nuevamente."
                        })
                st.rerun()
    
    with tab2:
        st.header("📖 Generador de Sesión de Aprendizaje")
        
        # Inicializar chat desde el inicio
        if 'chat_mensajes_sesion' not in st.session_state:
            st.session_state['chat_mensajes_sesion'] = []
        
        # Información contextual
        with st.expander("ℹ️ Información sobre la Sesión de Aprendizaje"):
            st.markdown("""
            **¿Qué incluye una sesión de aprendizaje?**
            - ✅ Competencias y capacidades a desarrollar
            - ✅ Secuencia didáctica (inicio, desarrollo, cierre)
            - ✅ Actividades de aprendizaje
            - ✅ Materiales y recursos
            - ✅ Evaluación formativa
            
            **Basado en:** Currículo Nacional de Educación Básica - MINEDU Perú
            """)
        
        # Mostrar información de unidad generada si existe
        if 'unidad_generada' in st.session_state:
            unidad_info = st.session_state['unidad_generada']
            st.info(f"📚 Unidad generada: {unidad_info.get('titulo', 'N/A')}")
            # Mostrar títulos de sesiones disponibles si existen
            titulos_sesiones = unidad_info.get('titulos_sesiones', [])
            if titulos_sesiones:
                st.success(f"📋 Se encontraron {len(titulos_sesiones)} sesiones en la unidad:")
                for idx, titulo_ses in enumerate(titulos_sesiones[:6], 1):  # Mostrar máximo 6
                    st.text(f"  {idx}. {titulo_ses}")
        
        with st.form("form_sesion", clear_on_submit=False):
            col1, col2 = st.columns(2)
            
            with col1:
                # Título de la unidad (viene de lo que se generó antes)
                titulo_unidad = st.text_input(
                    "📚 Título de la Unidad",
                    value=st.session_state.get('unidad_generada', {}).get('titulo', ''),
                    placeholder="Ej: La Materia y sus Propiedades",
                    help="Título de la unidad didáctica generada anteriormente (o ingresar manualmente)"
                )
                
                # Título de la sesión - mostrar selector si hay sesiones disponibles
                titulos_sesiones_disponibles = st.session_state.get('unidad_generada', {}).get('titulos_sesiones', [])
                if titulos_sesiones_disponibles:
                    titulo_sesion_seleccionado = st.selectbox(
                        "🎯 Título de la Sesión (selecciona de la unidad)",
                        options=[""] + titulos_sesiones_disponibles,
                        help="Puedes seleccionar una sesión de la unidad o escribir un título nuevo abajo"
                    )
                    titulo_sesion_personalizado = st.text_input(
                        "O escribe un título personalizado",
                        value="",
                        placeholder="Ej: Identificamos las propiedades de la materia",
                        help="Título específico de la sesión (se usará este si está lleno, o el seleccionado arriba)"
                    )
                    # Determinar qué título usar: personalizado tiene prioridad
                    titulo_sesion = titulo_sesion_personalizado.strip() if titulo_sesion_personalizado.strip() else (titulo_sesion_seleccionado if titulo_sesion_seleccionado else "")
                else:
                    titulo_sesion = st.text_input(
                        "🎯 Título de la Sesión",
                        placeholder="Ej: Identificamos las propiedades de la materia",
                        help="Título específico de la sesión de aprendizaje"
                    )
                
                # Competencia (viene de la unidad generada o se puede modificar/ingresar)
                competencias_extraidas = None
                if 'unidad_generada' in st.session_state:
                    contenido_unidad = st.session_state['unidad_generada'].get('contenido') or st.session_state.get('documento_raw_unidad', '')
                    if contenido_unidad:
                        competencias_extraidas = extraer_competencias_unidad_didactica(contenido_unidad)
                competencia = st.text_area(
                    "📋 Competencia",
                    value=competencias_extraidas or "",
                    placeholder="Competencia(s), capacidades y criterios de evaluación. Se llena automáticamente si generaste una unidad.",
                    help="Competencia a desarrollar. Se obtiene de la unidad didáctica generada o puedes ingresarla manualmente.",
                    height=80
                )
                
                # Tema
                tema = st.text_input(
                    "📝 Tema",
                    placeholder="Ej: Propiedades de la materia, Ecuaciones lineales, Comprensión de textos",
                    help="Tema o contenido específico de la sesión"
                )
            
            with col2:
                # Nivel fijo: Solo Secundaria
                nivel = "Secundaria"
                st.text_input(
                    "📊 Nivel",
                    value=nivel,
                    disabled=True,
                    help="Nivel educativo (limitado a Secundaria)"
                )
                
                # Mostrar grado de la unidad generada (si existe)
                grado_unidad = st.session_state.get('unidad_generada', {}).get('grado', '')
                if grado_unidad:
                    st.text_input(
                        "🎓 Grado",
                        value=grado_unidad,
                        disabled=True,
                        help="Grado obtenido de la unidad didáctica generada"
                    )
                    grado = grado_unidad
                else:
                    st.info("ℹ️ Genera primero una unidad didáctica para obtener el grado")
                    grado = ""
                
                seccion = st.text_input(
                    "👥 Sección",
                    placeholder="Ej: A, B, C",
                    help="Sección del grado"
                )
                
                duracion = st.text_input(
                    "⏱️ Duración",
                    placeholder="Ej: 90 minutos, 2 horas",
                    help="Duración de la sesión de aprendizaje"
                )
                
                # Metodología
                METODOLOGIAS = [
                    "— Seleccione metodología —",
                    "Aprendizaje Basado en Proyectos (ABP)",
                    "Aprendizaje Basado en Problemas",
                    "Flipped Classroom (Aula invertida)",
                    "Gamificación",
                    "Aprendizaje Cooperativo",
                    "Enfoque de Indagación Científica",
                    "Método de Casos",
                    "Design Thinking",
                    "Otro (describir en tema)"
                ]
                metodologia = st.selectbox(
                    "📐 Metodología",
                    options=METODOLOGIAS,
                    help="Metodología o enfoque pedagógico a utilizar en la sesión"
                )
            
            generar = st.form_submit_button("🎯 Generar Sesión de Aprendizaje", use_container_width=True)
        
        # FUERA del formulario - manejar resultados
        if generar:
            # Asegurar que todos los campos sean strings
            titulo_unidad = str(titulo_unidad) if titulo_unidad else ''
            titulo_sesion = str(titulo_sesion) if titulo_sesion else ''
            seccion = str(seccion) if seccion else ''
            duracion = str(duracion) if duracion else ''
            competencia = str(competencia) if competencia else ''
            tema = str(tema) if tema else ''
            metodologia = str(metodologia) if metodologia else ''
            
            # Obtener grado de la unidad generada si no está en el formulario
            grado_actual = str(grado) if 'grado' in locals() and grado else ''
            if not grado_actual and 'unidad_generada' in st.session_state:
                grado_actual = str(st.session_state['unidad_generada'].get('grado', '') or '')
            
            # Usar competencia del formulario; si está vacía, intentar extraer de la unidad
            competencias_para_sesion = competencia.strip() if competencia.strip() else None
            if not competencias_para_sesion and 'unidad_generada' in st.session_state:
                contenido_unidad = st.session_state['unidad_generada'].get('contenido') or st.session_state.get('documento_raw_unidad', '')
                if contenido_unidad:
                    competencias_para_sesion = extraer_competencias_unidad_didactica(contenido_unidad)
            
            # Metodología: no enviar si es el placeholder
            metodologia_para_sesion = metodologia.strip() if metodologia.strip() and not metodologia.startswith("—") else None
            
            # Validar campos requeridos
            if not titulo_unidad.strip():
                st.warning("⚠️ Por favor ingresa el título de la unidad")
            elif not titulo_sesion.strip():
                st.warning("⚠️ Por favor ingresa el título de la sesión")
            elif not nivel:
                st.warning("⚠️ Por favor selecciona el nivel")
            elif not grado_actual.strip():
                st.warning("⚠️ Por favor genera primero una unidad didáctica para obtener el grado")
            elif not seccion.strip():
                st.warning("⚠️ Por favor ingresa la sección")
            elif not duracion.strip():
                st.warning("⚠️ Por favor ingresa la duración")
            else:
                with st.spinner('🔄 Generando sesión de aprendizaje...'):
                    try:
                        resultado_raw = generar_sesion_aprendizaje(
                            titulo_unidad,
                            titulo_sesion,
                            nivel,
                            grado_actual,
                            seccion,
                            duracion,
                            competencias_unidad=competencias_para_sesion,
                            tema=tema.strip() or None,
                            metodologia=metodologia_para_sesion
                        )
                        
                        # Formatear el contenido con encabezados y estructura completa (incluye tabla I. DATOS INFORMATIVOS)
                        area_sesion = st.session_state.get('unidad_generada', {}).get('area_curricular', '') or ''
                        contenido_formateado = formatear_sesion_aprendizaje(
                            resultado_raw,
                            titulo_unidad,
                            titulo_sesion,
                            nivel,
                            grado_actual,
                            seccion,
                            duracion,
                            area_curricular=area_sesion
                        )
                        
                        # Guardar archivos automáticamente en Desktop
                        fecha_str = datetime.now().strftime('%Y%m%d_%H%M%S')
                        nombre_txt = f"sesion_aprendizaje_{fecha_str}.txt"
                        ruta_txt = guardar_archivo_desktop(contenido_formateado, nombre_txt, es_bytes=False)
                        
                        if DOCX_OK:
                            # Crear documento con título de unidad y título de sesión
                            doc_bytes = crear_documento_sesion_aprendizaje(
                                resultado_raw,
                                titulo_unidad,
                                titulo_sesion,
                                nivel,
                                grado_actual,
                                seccion,
                                duracion=duracion,
                                area_curricular=area_sesion
                            )
                            if doc_bytes:
                                nombre_docx = f"sesion_aprendizaje_{fecha_str}.docx"
                                ruta_docx = guardar_archivo_desktop(doc_bytes, nombre_docx, es_bytes=True)
                        else:
                            doc_bytes = None
                            ruta_docx = None
                        
                        # Guardar documento editable y metadatos para chat y descarga
                        st.session_state['documento_editable_sesion'] = contenido_formateado
                        st.session_state['documento_raw_sesion'] = resultado_raw
                        st.session_state['chat_mensajes_sesion'] = []
                        st.session_state['sesion_meta'] = {
                            'titulo_unidad': titulo_unidad,
                            'titulo_sesion': titulo_sesion,
                            'nivel': nivel,
                            'grado': grado_actual,
                            'seccion': seccion,
                            'duracion': duracion,
                            'area_curricular': area_sesion
                        }
                        
                        st.success("✅ ¡Sesión de aprendizaje generada exitosamente!")
                        if ruta_txt:
                            st.info(f"📁 Archivos guardados en: {ruta_txt.rsplit('/', 1)[0]}")
                                
                    except Exception as e:
                        st.error(f"❌ Error generando sesión de aprendizaje: {str(e)}")
                        st.info("💡 Verifica la conexión con AWS Bedrock")
        
        # Mostrar documento actual (generado o mejorado) y chat de mejoras - Sesión
        if st.session_state.get('documento_editable_sesion'):
            st.markdown("---")
            
            # Sección de documento con expander para mejor organización
            with st.expander("📄 Ver documento actual", expanded=True):
                doc_actual_sesion = st.session_state['documento_editable_sesion']
                st.markdown(doc_actual_sesion, unsafe_allow_html=True)
            
            # Botones de acción en una fila organizada
            st.markdown("### 📥 Descargar documento")
            col1, col2, col3 = st.columns([2, 2, 2])
            with col1:
                st.download_button(
                    "📄 Descargar TXT",
                    data=doc_actual_sesion,
                    file_name=f"sesion_aprendizaje_{datetime.now().strftime('%Y%m%d_%H%M%S')}.txt",
                    mime="text/plain",
                    key="download_txt_sesion",
                    use_container_width=True
                )
            with col2:
                if DOCX_OK:
                    meta = st.session_state.get('sesion_meta', {})
                    doc_bytes_sesion = crear_documento_sesion_aprendizaje(
                        st.session_state.get('documento_raw_sesion', doc_actual_sesion),
                        meta.get('titulo_unidad', ''),
                        meta.get('titulo_sesion', ''),
                        meta.get('nivel', 'Secundaria'),
                        meta.get('grado', ''),
                        meta.get('seccion', ''),
                        duracion=meta.get('duracion', ''),
                        area_curricular=meta.get('area_curricular', '')
                    )
                    if doc_bytes_sesion:
                        st.download_button(
                            "📝 Descargar WORD",
                            data=doc_bytes_sesion,
                            file_name=f"sesion_aprendizaje_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx",
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                            key="download_docx_sesion",
                            use_container_width=True
                        )
                    else:
                        st.button("📝 WORD no disponible", disabled=True, key="docx_disabled_sesion", use_container_width=True)
                else:
                    st.button("📝 WORD no disponible", disabled=True, key="docx_disabled_sesion", use_container_width=True)
            with col3:
                if st.button("🔄 Generar Nueva Sesión", key="nueva_sesion", use_container_width=True):
                    for k in ('documento_editable_sesion', 'documento_raw_sesion', 'chat_mensajes_sesion', 'sesion_meta'):
                        if k in st.session_state:
                            del st.session_state[k]
                    st.rerun()
            
            st.markdown("---")
        
        # Chat siempre visible desde el inicio - Mejor organizado
        st.markdown("### 💬 Editor de Chat - Mejorar sesión")
        
        # Mostrar estado del documento
        if st.session_state.get('documento_editable_sesion'):
            st.success("✅ Tienes una sesión generada. Puedes mejorarla usando el chat.")
        else:
            st.info("ℹ️ **Genera primero una sesión de aprendizaje arriba para poder mejorarla con el chat.**")
        
        st.info("💡 **Sugerencias:** Puedes pedir cambios como 'añade una actividad de cierre', 'simplifica las indicaciones', 'mejora la secuencia didáctica', etc.")
        
        # Contenedor para el chat con mejor estilo
        chat_container_sesion = st.container()
        with chat_container_sesion:
            # Mostrar historial de chat con mejor formato
            if st.session_state['chat_mensajes_sesion']:
                st.markdown("#### 📜 Historial de conversación")
                for idx, msg in enumerate(st.session_state['chat_mensajes_sesion']):
                    with st.chat_message(msg["role"]):
                        if msg["role"] == "user":
                            st.markdown(f"**Tu solicitud:**\n{msg['content']}")
                        else:
                            # Mejorar formato de respuesta del asistente
                            contenido = msg['content']
                            if "✅ Cambios aplicados" in contenido:
                                st.success("✅ **Cambios aplicados exitosamente**")
                                # Extraer solo la vista previa si existe
                                if "Vista previa:" in contenido:
                                    partes = contenido.split("Vista previa:", 1)
                                    if len(partes) > 1:
                                        st.markdown(f"**Vista previa:**\n{partes[1].strip()}")
                            elif "❌ Error" in contenido:
                                st.error(contenido)
                            else:
                                st.markdown(contenido)
                    if idx < len(st.session_state['chat_mensajes_sesion']) - 1:
                        st.markdown("---")
            else:
                st.markdown("*No hay mensajes aún. Escribe abajo para comenzar a mejorar la sesión.*")
            
            # Input de chat con mejor placeholder (habilitado solo si hay documento)
            tiene_documento_sesion = st.session_state.get('documento_editable_sesion')
            prompt_sesion = st.chat_input(
                "Escribe aquí cómo quieres mejorar la sesión..." if tiene_documento_sesion else "Primero genera una sesión de aprendizaje arriba...",
                key="chat_input_sesion",
                disabled=not tiene_documento_sesion
            )
            
            if prompt_sesion and tiene_documento_sesion:
                # Agregar mensaje del usuario al historial
                st.session_state['chat_mensajes_sesion'].append({"role": "user", "content": prompt_sesion})
                
                # Mostrar spinner mientras se procesa
                with st.spinner("🔄 Aplicando cambios a la sesión..."):
                    try:
                        nuevo_doc_sesion = mejorar_documento_con_instruccion(
                            st.session_state['documento_editable_sesion'],
                            prompt_sesion,
                            "sesión de aprendizaje"
                        )
                        
                        if nuevo_doc_sesion and not nuevo_doc_sesion.startswith("[Error"):
                            st.session_state['documento_editable_sesion'] = nuevo_doc_sesion
                            st.session_state['documento_raw_sesion'] = nuevo_doc_sesion
                            # Mensaje de éxito más claro
                            st.session_state['chat_mensajes_sesion'].append({
                                "role": "assistant",
                                "content": f"✅ Cambios aplicados exitosamente. La sesión ha sido actualizada.\n\n**Vista previa de los cambios:**\n\n{nuevo_doc_sesion[:400]}..."
                            })
                        else:
                            st.session_state['chat_mensajes_sesion'].append({
                                "role": "assistant",
                                "content": f"⚠️ No se pudieron aplicar los cambios. Por favor, intenta con una instrucción más específica."
                            })
                    except Exception as e:
                        st.session_state['chat_mensajes_sesion'].append({
                            "role": "assistant",
                            "content": f"❌ Error al procesar la solicitud: {str(e)}\n\nPor favor, intenta nuevamente."
                        })
                st.rerun()

else:
    st.error("⚠️ Los servicios no están disponibles. Verifica la configuración.")
    
    with st.expander("🔧 Información de diagnóstico"):
        st.write(f"**Archivo actual:** {__file__}")
        st.write(f"**Directorio actual:** {os.getcwd()}")
        st.write(f"**Directorio del archivo:** {os.path.dirname(__file__)}")
        st.write(f"**Directorio padre:** {os.path.dirname(os.path.dirname(__file__))}")
        
        core_path = os.path.join(os.path.dirname(os.path.dirname(__file__)), 'core')
        st.write(f"**Buscando core en:** {core_path}")
        st.write(f"**Core existe:** {os.path.exists(core_path)}")
        
        if os.path.exists(core_path):
            st.write(f"**Archivos en core:** {os.listdir(core_path)}")

# Footer
st.markdown("---")
st.markdown("""
<div style='text-align: center'>
<h6>🎓 Sistema de Generación de Contenido Educativo con IA</h6>
<p><em>Desarrollado para el Ministerio de Educación del Perú</em></p>
</div>
""", unsafe_allow_html=True)